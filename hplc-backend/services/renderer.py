from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from models import InspectionReport, ValidationReport, InspectionItem, ValidationSection
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os

# 注册中文字体（使用系统字体）
def _register_chinese_font():
    font_paths = [
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\msyh.ttc",
        r"/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
    ]
    for path in font_paths:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont("Chinese", path))
                return "Chinese"
            except Exception:
                continue
    return "Helvetica"

_CN_FONT = _register_chinese_font()

def _pdf_styles():
    styles = getSampleStyleSheet()
    base = {"fontName": _CN_FONT, "fontSize": 10, "leading": 16}
    return {
        "h1": ParagraphStyle("h1", fontName=_CN_FONT, fontSize=14, leading=20, alignment=1, spaceAfter=6),
        "h2": ParagraphStyle("h2", fontName=_CN_FONT, fontSize=12, leading=18, spaceBefore=8, spaceAfter=4),
        "body": ParagraphStyle("body", **base, spaceAfter=4),
    }

def _pdf_table(headers, rows, col_widths=None):
    data = [[Paragraph(str(h), ParagraphStyle("th", fontName=_CN_FONT, fontSize=9, leading=14)) for h in headers]]
    for r in rows:
        data.append([Paragraph(str(v), ParagraphStyle("td", fontName=_CN_FONT, fontSize=9, leading=14)) for v in r])
    t = Table(data, colWidths=col_widths)
    t.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.9, 0.9, 0.9)),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    return t

# ── helpers ───────────────────────────────────────────────────────────────────

def _heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    return p

def _para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(24)
    return p

def _cell_text(cell, text, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def _set_document_fonts(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(12)

def _rs_rows(raw):
    return raw if isinstance(raw, list) else []

def _render_rs_table(doc, title, headers, rows, merged_columns=()):
    _heading(doc, title, level=2)
    rows = _rs_rows(rows)
    table = doc.add_table(rows=1 + max(1, len(rows)), cols=len(headers))
    table.style = "Table Grid"
    _table_header_row(table, headers)
    for ri, values in enumerate(rows):
        for ci, value in enumerate(values[:len(headers)]):
            _cell_text(table.rows[ri + 1].cells[ci], str(value), size=10.5)
    if len(rows) > 1:
        for col in merged_columns:
            if col < len(headers):
                first = table.rows[1].cells[col]
                for ri in range(2, len(rows) + 1):
                    first = first.merge(table.rows[ri].cells[col])
    return table

def _render_linear_chart(doc, title, values):
    if not isinstance(values, list) or len(values) < 2:
        return
    try:
        x_values = [float(value) for value in values[0][1:] if str(value).strip()]
        y_values = [float(value) for value in values[1][1:] if str(value).strip()]
        count = min(len(x_values), len(y_values))
        if count < 2:
            return
        x_values, y_values = x_values[:count], y_values[:count]
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import numpy as np
        slope, intercept = np.polyfit(x_values, y_values, 1)
        figure, axis = plt.subplots(figsize=(6.4, 3.8), dpi=160)
        axis.scatter(x_values, y_values, color="#304f9e", s=34, zorder=3, label="Observed")
        x_line = np.linspace(min(x_values), max(x_values), 100)
        axis.plot(x_line, slope * x_line + intercept, color="#d97706", linewidth=1.8, label="Linear fit")
        r_squared = values[3][1] if len(values) > 3 and len(values[3]) > 1 else ""
        axis.set_xlabel("Concentration (ug/ml)")
        axis.set_ylabel("Peak area")
        axis.grid(alpha=0.25)
        axis.legend(loc="best", frameon=False)
        axis.text(0.02, 0.98, f"y = {slope:.4f}x {'+' if intercept >= 0 else '-'} {abs(intercept):.4f}\nR² = {r_squared}", transform=axis.transAxes, va="top", fontsize=9)
        figure.tight_layout()
        image = BytesIO()
        figure.savefig(image, format="png")
        plt.close(figure)
        image.seek(0)
        doc.add_picture(image, width=Cm(15))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        return

def _render_related_substances(doc, data, is_assay=False):
    section_1 = [("1.1 仪器", "instruments"), ("1.2 试剂", "reagents"), ("1.3 对照品与样品", "references")]
    headers = {
        "instruments": ["名称", "型号", "编号", "厂家", "有效期"], "reagents": ["名称", "批号", "级别", "厂家"],
        "references": ["名称", "批号", "含量", "厂家"], "compounds": ["化合物名称", "纯度%", "称样量mg", "定容体积ml", "稀释剂", "母液浓度mg/ml"],
        "single_standards": ["化合物名称", "母液浓度mg/ml", "移取体积ml", "定容体积ml", "稀释剂", "最终浓度ug/ml"],
        "linear_a_stock": ["杂质名称", "母液浓度mg/ml", "移取体积ml", "定容体积ml", "稀释剂", "线性储备液浓度ug/ml"],
        "linear_b_stock": ["杂质名称", "称样量mg", "纯度%", "定容体积ml", "稀释剂", "移取量ml", "定容体积ml", "稀释剂", "线性储备液b浓度ug/ml"],
        "mobile_phase": ["名称", "称取/量取", "定容体积ml", "调节/稀释", "批号", "备注"], "blank": ["名称", "组成/来源", "批号", "备注"],
        "system_suitability": ["化合物名称", "母液浓度mg/ml", "移取体积ml", "定容体积ml", "稀释剂", "最终浓度ug/ml"],
        "samples": ["化合物名称", "批号", "称样量mg", "定容体积ml", "稀释剂", "最终浓度mg/ml"],
        "reference_solution": ["化合物名称", "浓度mg/ml", "移取量ml", "定容体积ml", "最终浓度ug/ml"],
        "liquid_method": ["项目", "条件", "确定依据"], "response_factors": ["名称", "a-线性斜率", "a-相对响应因子", "b-线性斜率", "b-相对响应因子", "相对响应因子均值"],
        "suitability_results": ["系统适用性", "保留时间min", "峰面积", "USP-分离度", "理论塔板数"],
    }
    for title, key in section_1:
        _render_rs_table(doc, title, headers[key], data.get(key, []))
    _render_rs_table(doc, "2.1 各标准品母液及单标配置", headers["compounds"], data.get("compounds", []), (4,))
    if not is_assay:
        _render_rs_table(doc, "2.2 各定位溶液单标配置", headers["single_standards"], data.get("single_standards", []), (3, 4))
    if not is_assay or data.get("assay_stock_enabled"):
        _render_rs_table(doc, "2.2 含量线性储备液配置" if is_assay else "2.3 线性储备液a配置", headers["linear_a_stock"], data.get("linear_a_stock", []), (3, 4))
    if not is_assay:
        _render_rs_table(doc, "2.4 线性储备液b配置", headers["linear_b_stock"], data.get("linear_b_stock", []), (3, 4, 5, 6, 7))
    linear_sections = [("2.3 线性储备液线性配置", "linear_a")] if is_assay else [("2.5 线性储备液a线性配置", "linear_a"), ("2.6 线性储备液b线性配置", "linear_b")]
    for label, key in linear_sections:
        for index, values in enumerate(_rs_rows(data.get(key))):
            _render_rs_table(doc, f"{label} {index + 1}", ["名称", "储备液浓度ug/ml", "移取体积ml", "定容体积ml", "线性点浓度ug/ml", "稀释剂"], values, (1, 5))
    step_number = 4 if is_assay else 7
    _render_rs_table(doc, f"2.{step_number} 流动相配制", headers["mobile_phase"], data.get("mobile_phase", []))
    _render_rs_table(doc, f"2.{step_number + 1} 空白溶液", headers["blank"], data.get("blank", []))
    _render_rs_table(doc, f"2.{step_number + 2} 系统适用性溶液", headers["system_suitability"], data.get("system_suitability", []), (4,))
    sample_headers = data.get("assay_samples_headers") if is_assay else headers["samples"]
    sample_merge_columns = (6,) if is_assay else (4,)
    _render_rs_table(doc, f"2.{step_number + 3} 供试品溶液浓度", sample_headers or headers["samples"], data.get("samples", []), sample_merge_columns)
    _render_rs_table(doc, f"2.{step_number + 4} 对照溶液", data.get("reference_solution_headers") or headers["reference_solution"], data.get("reference_solution", []))
    if data.get("reference_solution_notes"):
        _para(doc, f"备注：{data['reference_solution_notes']}")
    _render_rs_table(doc, "3.1 液相方法", headers["liquid_method"], data.get("liquid_method", []))
    for index, method in enumerate(_rs_rows(data.get("hplc_methods"))):
        condition_rows = [
            ["色谱柱", method.get("column", "")], ["柱温(℃)", method.get("temp", "")],
            ["进样盘温度(℃)", method.get("tray_temp", "")], ["进样量(ul)", method.get("injection_volume", "")],
            ["流速(ml/min)", method.get("flow_rate", "")], ["流动相A", method.get("mobile_a", "")],
            ["流动相B", method.get("mobile_b", "")], ["切换阀", method.get("switch_valve", "")],
        ]
        _render_rs_table(doc, f"3.1 液相方法{index + 1}", ["项目", "条件"], condition_rows)
        _render_rs_table(doc, "梯度洗脱", method.get("gradient_headers") or ["时间min", "流动相A%", "流动相B%"], method.get("gradient", []))
        if method.get("notes"):
            _para(doc, f"备注：{method['notes']}")
    analysis_sections = ([("4.1 含量线性分析", "linear_a_analysis")] if data.get("assay_linearity_enabled") else []) if is_assay else [("4.1 线性a分析", "linear_a_analysis"), ("4.2 线性b分析", "linear_b_analysis")]
    for label, key in analysis_sections:
        for index, values in enumerate(_rs_rows(data.get(key))):
            _render_rs_table(doc, f"{label} {index + 1}", ["名称", "s1", "s2", "s3", "s4", "s5"], values)
            _render_linear_chart(doc, f"{label} {index + 1}", values)
    if is_assay:
        _render_rs_table(doc, f"4.2 {data.get('assay_reference_title') or '对照峰面积'}", data.get("peak_area_headers") or ["名称", "浓度ug/ml", "峰面积1", "峰面积均值"], data.get("peak_areas", []))
        _render_rs_table(doc, f"4.3 {data.get('assay_calculation_title') or '含量计算'}", data.get("content_headers") or ["名称", "浓度ug/ml", "峰面积1", "峰面积均值"], data.get("contents", []))
        return
    _render_rs_table(doc, "4.3 相对响应因子", headers["response_factors"], data.get("response_factors", []))
    _render_rs_table(doc, "4.4.1 系统适用性结果", headers["suitability_results"], data.get("suitability_results", []))
    for title, key in [("4.4.2 峰面积", "peak_areas"), ("4.4.3 含量", "contents")]:
        rows = _rs_rows(data.get(key))
        header_key = "peak_area_headers" if key == "peak_areas" else "content_headers"
        headers_for_table = data.get(header_key) or ["批号", "杂质1", "杂质2", "主峰"]
        _render_rs_table(doc, title, headers_for_table, rows)

def _table_header_row(table, headers):
    row = table.rows[0]
    for i, h in enumerate(headers):
        _cell_text(row.cells[i], h, bold=True)

def _add_kv_table(doc, pairs):
    """2-col key-value table"""
    t = doc.add_table(rows=len(pairs), cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(pairs):
        _cell_text(t.rows[i].cells[0], k, bold=True)
        _cell_text(t.rows[i].cells[1], v)
    return t

def _conclusion_row(table):
    row = table.add_row()
    row.cells[0].merge(row.cells[-1])
    _cell_text(row.cells[0], "□符合规定     □不符合规定")

# ── inspection item block ─────────────────────────────────────────────────────

def _render_item(doc, item: InspectionItem, index: int):
    _heading(doc, f"{'检查' if index > 0 else ''}{index + 1}. {item.title}", level=2)

    # temp/humidity
    p = doc.add_paragraph()
    p.add_run(f"温度（℃）：{item.temperature}    相对湿度（%）：{item.humidity}")

    # instruments
    if item.instruments:
        headers = ["仪器", "型号", "编号", "厂家", "有效期"]
        t = doc.add_table(rows=1 + len(item.instruments), cols=len(headers))
        t.style = "Table Grid"
        _table_header_row(t, headers)
        for ri, row_data in enumerate(item.instruments):
            for ci, val in enumerate(row_data[:len(headers)]):
                _cell_text(t.rows[ri + 1].cells[ci], val)

    # reagents
    if item.reagents:
        headers = ["试剂/对照品", "批号", "级别/含量", "来源"]
        t = doc.add_table(rows=1 + len(item.reagents), cols=len(headers))
        t.style = "Table Grid"
        _table_header_row(t, headers)
        for ri, row_data in enumerate(item.reagents):
            for ci, val in enumerate(row_data[:len(headers)]):
                _cell_text(t.rows[ri + 1].cells[ci], val)

    for label, text in [
        ("标准", item.standard),
        ("检验方法", item.method),
        ("操作过程", item.procedure),
        ("计算过程及结果", item.calculation),
        ("结果", item.result),
    ]:
        if text:
            p = doc.add_paragraph()
            p.add_run(f"{label}：").bold = True
            p.add_run(text)

    # conclusion / signatures
    t = doc.add_table(rows=3, cols=4)
    t.style = "Table Grid"
    _cell_text(t.rows[0].cells[0], "结论", bold=True)
    t.rows[0].cells[1].merge(t.rows[0].cells[3])
    _cell_text(t.rows[0].cells[1], item.conclusion)
    _cell_text(t.rows[1].cells[0], "检验人")
    _cell_text(t.rows[1].cells[1], item.inspector)
    _cell_text(t.rows[1].cells[2], "检验日期")
    _cell_text(t.rows[1].cells[3], item.inspector_date)
    _cell_text(t.rows[2].cells[0], "复核人")
    _cell_text(t.rows[2].cells[1], item.reviewer)
    _cell_text(t.rows[2].cells[2], "复核日期")
    _cell_text(t.rows[2].cells[3], item.reviewer_date)

# ── inspection report ─────────────────────────────────────────────────────────

def render_inspection_word(data: InspectionReport) -> bytes:
    doc = Document()
    _set_document_fonts(doc)

    pairs = [
        ("样品名称", data.product_name),
        ("样品来源", data.source),
        ("项目类别", data.sample_nature),
        ("方法提供", data.basis),
        ("项目内容", data.request_date),
    ]
    _add_kv_table(doc, pairs)

    # approvals
    if data.approvals:
        _heading(doc, "实验人与复核人", level=2)
        t = doc.add_table(rows=1 + len(data.approvals), cols=3)
        t.style = "Table Grid"
        for i, h in enumerate(["职责/实验人", "复核人", "签名/日期"]):
            _cell_text(t.rows[0].cells[i], h, bold=True)
        for ri, row in enumerate(data.approvals):
            _cell_text(t.rows[ri + 1].cells[0], row.role)
            _cell_text(t.rows[ri + 1].cells[1], row.position)
            _cell_text(t.rows[ri + 1].cells[2], row.signature_date)

    # inspection items
    for i, item in enumerate(data.items):
        doc.add_paragraph()
        if item.title == "有关物质" and item.related_substances:
            _heading(doc, "四、有关物质", level=2)
            _render_related_substances(doc, item.related_substances)
        elif item.title == "含量测定" and item.assay_data:
            _heading(doc, "五、含量", level=2)
            _render_related_substances(doc, item.assay_data, is_assay=True)
        else:
            _render_item(doc, item, i)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_inspection_html(data: InspectionReport) -> str:
    import html as _html
    e = _html.escape
    type_label = {"api": "原料药", "solid": "固体制剂", "liquid": "液体制剂"}.get(data.report_type, "")
    rows = ""
    for i, item in enumerate(data.items):
        inst_rows = "".join(f"<tr>{''.join(f'<td>{e(v)}</td>' for v in r)}</tr>" for r in item.instruments)
        reag_rows = "".join(f"<tr>{''.join(f'<td>{e(v)}</td>' for v in r)}</tr>" for r in item.reagents)
        rows += f"""
        <div class="item">
          <h3>{i+1}. {e(item.title)}</h3>
          <p>温度：{e(item.temperature)} ℃　相对湿度：{e(item.humidity)} %</p>
          {f'<table><tr><th>仪器</th><th>型号</th><th>编号</th><th>厂家</th><th>有效期</th></tr>{inst_rows}</table>' if item.instruments else ''}
          {f'<table><tr><th>试剂/对照品</th><th>批号</th><th>级别/含量</th><th>来源</th></tr>{reag_rows}</table>' if item.reagents else ''}
          <p><b>标准：</b>{e(item.standard)}</p>
          <p><b>操作过程：</b>{e(item.procedure)}</p>
          <p><b>结果：</b>{e(item.result)}</p>
          <p><b>结论：</b>{e(item.conclusion)}</p>
        </div>"""

    return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
    <style>body{{font-family:SimSun,serif;padding:20px}}
    table{{border-collapse:collapse;width:100%;margin:8px 0}}
    th,td{{border:1px solid #999;padding:4px 8px;font-size:13px}}
    th{{background:#eee}}.item{{border:1px solid #ddd;padding:12px;margin:12px 0}}
    h1,h2{{text-align:center}}h3{{color:#333}}</style></head>
    <body>
    <h1>{e(data.product_name)}</h1>
    <h2>{e(type_label)}检验记录　{e(data.doc_number)}</h2>
    <p>批号：{e(data.batch_number)}　规格：{e(data.specification)}　来源：{e(data.source)}</p>
    <p>检验依据：{e(data.basis)}</p>
    {rows}
    <p><b>异常情况说明：</b>{e(data.abnormal)}</p>
    </body></html>"""


# ── validation section ────────────────────────────────────────────────────────

def _render_validation_section(doc, title: str, section: ValidationSection):
    _heading(doc, title, level=2)
    for i, para in enumerate(section.paragraphs):
        _para(doc, f"（{i+1}）{para}")
    if section.table_headers and section.table_rows:
        t = doc.add_table(rows=1 + len(section.table_rows), cols=len(section.table_headers))
        t.style = "Table Grid"
        _table_header_row(t, section.table_headers)
        for ri, row_data in enumerate(section.table_rows):
            for ci, val in enumerate(row_data[:len(section.table_headers)]):
                _cell_text(t.rows[ri + 1].cells[ci], val)
    if section.conclusion:
        p = doc.add_paragraph()
        p.add_run("结论：").bold = True
        p.add_run(section.conclusion)


def render_validation_word(data: ValidationReport) -> bytes:
    doc = Document()
    _heading(doc, data.project_name or "方法学验证报告")
    if data.doc_number:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"文件编号：{data.doc_number}")

    # 封面签名
    t = doc.add_table(rows=3, cols=3)
    t.style = "Table Grid"
    for i, (role, person, date) in enumerate([
        ("撰写人", data.author, data.author_date),
        ("审核人", data.reviewer, data.reviewer_date),
        ("批准人", data.approver, data.approver_date),
    ]):
        _cell_text(t.rows[i].cells[0], role, bold=True)
        _cell_text(t.rows[i].cells[1], person)
        _cell_text(t.rows[i].cells[2], date)

    overview = [
        ("验证目的", data.purpose),
        ("适用范围", data.scope),
        ("验证依据", data.basis),
        ("方法概要", data.method_summary),
    ]
    if any(v for _, v in overview):
        _heading(doc, "1. 概述", level=2)
        for label, text in overview:
            if text:
                p = doc.add_paragraph()
                p.add_run(f"{label}：").bold = True
                p.add_run(text)

    # 检测条件
    _heading(doc, "2. 检测条件", level=2)
    cond = data.conditions
    cond_rows = [
        ("色谱柱", cond.column.condition, cond.column.basis),
        ("柱温", cond.temp.condition, cond.temp.basis),
        ("流动相", cond.mobile_phase.condition, cond.mobile_phase.basis),
        ("波长", cond.wavelength.condition, cond.wavelength.basis),
        ("流速", cond.flow_rate.condition, cond.flow_rate.basis),
        ("进样体积", cond.injection_vol.condition, cond.injection_vol.basis),
        ("限度", cond.limit.condition, cond.limit.basis),
        ("运行时间", cond.run_time.condition, cond.run_time.basis),
    ]
    t = doc.add_table(rows=1 + len(cond_rows), cols=3)
    t.style = "Table Grid"
    for i, h in enumerate(["检测条件", "参数", "确定依据"]):
        _cell_text(t.rows[0].cells[i], h, bold=True)
    for ri, (label, val, basis) in enumerate(cond_rows):
        _cell_text(t.rows[ri + 1].cells[0], label)
        _cell_text(t.rows[ri + 1].cells[1], val)
        _cell_text(t.rows[ri + 1].cells[2], basis)

    if data.validation_summary:
        _heading(doc, "3. 验证项目与可接受标准", level=2)
        t = doc.add_table(rows=1 + len(data.validation_summary), cols=3)
        t.style = "Table Grid"
        _table_header_row(t, ["项目", "可接受标准", "结果"])
        for ri, row_data in enumerate(data.validation_summary):
            for ci, val in enumerate(row_data[:3]):
                _cell_text(t.rows[ri + 1].cells[ci], val)

    # 仪器与试剂
    _heading(doc, "4. 仪器与试剂", level=2)
    for title, headers, rows in [
        ("仪器", ["名称", "编号", "型号", "厂家"], data.instruments),
        ("试剂", ["名称", "批号", "级别", "厂家"], data.reagents),
        ("对照品", ["名称", "批号", "含量(%)", "厂家"], data.references),
        ("样品", ["名称", "批号", "厂家"], data.samples),
        ("辅料", ["名称", "批号", "厂家"], data.excipients),
    ]:
        if rows:
            _heading(doc, title, level=2)
            t = doc.add_table(rows=1 + len(rows), cols=len(headers))
            t.style = "Table Grid"
            _table_header_row(t, headers)
            for ri, row_data in enumerate(rows):
                for ci, val in enumerate(row_data[:len(headers)]):
                    _cell_text(t.rows[ri + 1].cells[ci], val)

    # 验证内容
    _heading(doc, "5. 验证内容")
    sections = [
        ("5.1 专属性", data.s51_specificity),
        ("5.2 线性", data.s52_linearity),
        ("5.3 定量限和检测限", data.s53_loq_lod),
        ("5.4 重复性", data.s54_repeatability),
        ("5.5 中间精密度", data.s55_intermediate),
        ("5.6 准确度", data.s56_accuracy),
        ("5.7 溶液稳定性", data.s57_stability),
        ("5.8 耐用性", data.s58_robustness),
    ]
    for title, section in sections:
        _render_validation_section(doc, title, section)

    # 修订历史
    if data.revisions:
        _heading(doc, "6. 修订历史", level=2)
        t = doc.add_table(rows=1 + len(data.revisions), cols=3)
        t.style = "Table Grid"
        _table_header_row(t, ["编号", "更改原因", "生效日期"])
        for ri, row_data in enumerate(data.revisions):
            for ci, val in enumerate(row_data[:3]):
                _cell_text(t.rows[ri + 1].cells[ci], val)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_inspection_pdf(data: InspectionReport) -> bytes:
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    s = _pdf_styles()
    type_label = {"api": "原料药", "solid": "固体制剂", "liquid": "液体制剂"}.get(data.report_type, "")
    story = [
        Paragraph(data.product_name or "检验记录", s["h1"]),
        Paragraph(f"{type_label}检验记录　{data.doc_number}", s["h2"]),
        _pdf_table(["检品名称", "批号", "规格", "来源", "检验依据"],
                   [[data.product_name, data.batch_number, data.specification, data.source, data.basis]]),
        Spacer(1, 0.3*cm),
    ]
    for i, item in enumerate(data.items):
        story.append(Paragraph(f"{i+1}. {item.title}", s["h2"]))
        story.append(Paragraph(f"温度：{item.temperature} ℃　相对湿度：{item.humidity} %", s["body"]))
        if item.instruments:
            story.append(_pdf_table(["仪器", "型号", "编号", "厂家", "有效期"], item.instruments))
        if item.reagents:
            story.append(_pdf_table(["试剂/对照品", "批号", "级别/含量", "来源"], item.reagents))
        for label, text in [("标准", item.standard), ("操作过程", item.procedure),
                             ("结果", item.result), ("结论", item.conclusion)]:
            if text:
                story.append(Paragraph(f"<b>{label}：</b>{text}", s["body"]))
        story.append(Spacer(1, 0.2*cm))
    if data.abnormal:
        story.append(Paragraph(f"<b>异常情况说明：</b>{data.abnormal}", s["body"]))
    doc.build(story)
    return buf.getvalue()


def render_validation_pdf(data: ValidationReport) -> bytes:
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    s = _pdf_styles()
    story = [
        Paragraph(data.project_name or "方法学验证报告", s["h1"]),
        Paragraph(f"文件编号：{data.doc_number}", s["body"]),
        Spacer(1, 0.3*cm),
    ]
    overview_rows = [["验证目的", data.purpose], ["适用范围", data.scope], ["验证依据", data.basis], ["方法概要", data.method_summary]]
    if any(row[1] for row in overview_rows):
        story.append(Paragraph("1. 概述", s["h2"]))
        story.append(_pdf_table(["项目", "内容"], [row for row in overview_rows if row[1]]))
    cond = data.conditions
    story.append(Paragraph("2. 检测条件", s["h2"]))
    story.append(_pdf_table(
        ["检测条件", "参数", "确定依据"],
        [["色谱柱", cond.column.condition, cond.column.basis],
         ["柱温", cond.temp.condition, cond.temp.basis],
         ["流动相", cond.mobile_phase.condition, cond.mobile_phase.basis],
         ["波长", cond.wavelength.condition, cond.wavelength.basis],
         ["流速", cond.flow_rate.condition, cond.flow_rate.basis],
         ["进样体积", cond.injection_vol.condition, cond.injection_vol.basis],
         ["限度", cond.limit.condition, cond.limit.basis],
         ["运行时间", cond.run_time.condition, cond.run_time.basis]],
    ))
    if data.validation_summary:
        story.append(Paragraph("3. 验证项目与可接受标准", s["h2"]))
        story.append(_pdf_table(["项目", "可接受标准", "结果"], data.validation_summary))
    sections = [
        ("5.1 专属性", data.s51_specificity), ("5.2 线性", data.s52_linearity),
        ("5.3 定量限和检测限", data.s53_loq_lod), ("5.4 重复性", data.s54_repeatability),
        ("5.5 中间精密度", data.s55_intermediate), ("5.6 准确度", data.s56_accuracy),
        ("5.7 溶液稳定性", data.s57_stability), ("5.8 耐用性", data.s58_robustness),
    ]
    for title, sec in sections:
        story.append(Paragraph(title, s["h2"]))
        for i, p in enumerate(sec.paragraphs):
            story.append(Paragraph(f"（{i+1}）{p}", s["body"]))
        if sec.table_headers and sec.table_rows:
            story.append(_pdf_table(sec.table_headers, sec.table_rows))
        if sec.conclusion:
            story.append(Paragraph(f"<b>结论：</b>{sec.conclusion}", s["body"]))
    doc.build(story)
    return buf.getvalue()


def render_validation_html(data: ValidationReport) -> str:
    import html as _html
    e = _html.escape

    def section_html(title, section: ValidationSection):
        paras = "".join(f"<p>（{i+1}）{e(p)}</p>" for i, p in enumerate(section.paragraphs))
        if section.table_headers and section.table_rows:
            heads = "".join(f"<th>{e(h)}</th>" for h in section.table_headers)
            body = "".join(f"<tr>{''.join(f'<td>{e(v)}</td>' for v in r)}</tr>" for r in section.table_rows)
            tbl = f"<table><tr>{heads}</tr>{body}</table>"
        else:
            tbl = ""
        concl = f"<p><b>结论：</b>{e(section.conclusion)}</p>" if section.conclusion else ""
        return f"<h3>{e(title)}</h3>{paras}{tbl}{concl}"

    sections_html = "".join(section_html(t, s) for t, s in [
        ("5.1 专属性", data.s51_specificity),
        ("5.2 线性", data.s52_linearity),
        ("5.3 定量限和检测限", data.s53_loq_lod),
        ("5.4 重复性", data.s54_repeatability),
        ("5.5 中间精密度", data.s55_intermediate),
        ("5.6 准确度", data.s56_accuracy),
        ("5.7 溶液稳定性", data.s57_stability),
        ("5.8 耐用性", data.s58_robustness),
    ])
    overview_rows = [["验证目的", data.purpose], ["适用范围", data.scope], ["验证依据", data.basis], ["方法概要", data.method_summary]]
    overview_html = _html_table(["项目", "内容"], [r for r in overview_rows if r[1]], e)
    summary_html = _html_table(["项目", "可接受标准", "结果"], data.validation_summary, e)

    return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
    <style>body{{font-family:SimSun,serif;padding:20px}}
    table{{border-collapse:collapse;width:100%;margin:8px 0}}
    th,td{{border:1px solid #999;padding:4px 8px;font-size:13px}}
    th{{background:#eee}}h1,h2{{text-align:center}}</style></head>
    <body>
    <h1>{e(data.project_name)}</h1>
    <p>文件编号：{e(data.doc_number)}</p>
    {f'<h2>1. 概述</h2>{overview_html}' if overview_html else ''}
    {f'<h2>3. 验证项目与可接受标准</h2>{summary_html}' if summary_html else ''}
    <h2>5. 验证内容</h2>
    {sections_html}
    </body></html>"""


def _html_table(headers, rows, esc):
    if not rows:
        return ""
    heads = "".join(f"<th>{esc(str(h))}</th>" for h in headers)
    body = "".join("<tr>" + "".join(f"<td>{esc(str(v))}</td>" for v in row[:len(headers)]) + "</tr>" for row in rows)
    return f"<table><tr>{heads}</tr>{body}</table>"
