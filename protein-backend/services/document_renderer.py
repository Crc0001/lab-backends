from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import io

TEMPLATE_PATH = Path("template.docx")


def _ensure_template():
    """若模板不存在则程序化生成一份标准模板。"""
    if TEMPLATE_PATH.exists():
        return

    doc = Document()

    # 标题
    title = doc.add_heading("蛋白发酵实验报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("报告编号：{{report_title}}")
    doc.add_paragraph("生成日期：{{created_at}}")
    doc.add_paragraph()

    # 一、材料与试剂
    doc.add_heading("一、材料与试剂", level=1)
    doc.add_paragraph("表达载体：{{vector_name}}")
    doc.add_paragraph("宿主菌株：{{host_cell}}")
    doc.add_paragraph("目的蛋白大小：{{target_gene_size_kda}} kDa")
    doc.add_paragraph("序列标签：{{tags}}")
    doc.add_paragraph()
    doc.add_paragraph("主要试剂：")
    doc.add_paragraph("{{reagents_table}}")
    doc.add_paragraph("主要仪器：")
    doc.add_paragraph("{{instruments_table}}")

    # 二、发酵诱导工艺
    doc.add_heading("二、发酵诱导工艺", level=1)
    doc.add_paragraph("培养基类型：{{media_type}}")
    doc.add_paragraph("接种量：{{inoculation_volume_percent}}%")
    doc.add_paragraph("培养温度：{{culture_temperature}} ℃")
    doc.add_paragraph("摇床转速：{{shaker_speed_rpm}} rpm")
    doc.add_paragraph("诱导时 OD600：{{induction_od600}}")
    doc.add_paragraph("IPTG 终浓度：{{iptg_concentration_mm}} mM")
    doc.add_paragraph("诱导温度：{{induction_temperature}} ℃")

    # 三、表达形式
    doc.add_heading("三、蛋白表达形式", level=1)
    doc.add_paragraph("表达形式：{{expression_form}}")
    doc.add_paragraph("包涵体处理工艺：{{inclusion_body_treatment}}")

    # 四、纯化结果
    doc.add_heading("四、纯化结果", level=1)
    doc.add_paragraph("纯化方法：{{purification_method}}")
    doc.add_paragraph("蛋白定量方法：{{assay_method}}")
    doc.add_paragraph("蛋白浓度：{{protein_concentration_mg_ml}} mg/mL")
    doc.add_paragraph("蛋白纯度：{{protein_purity_percent}}%")

    # 五、电泳图
    doc.add_heading("五、SDS-PAGE 电泳图", level=1)
    doc.add_paragraph("{{sds_page_image}}")

    # 六、质谱图
    doc.add_heading("六、质谱图", level=1)
    doc.add_paragraph("{{mass_spec_image}}")

    doc.save(str(TEMPLATE_PATH))


def _resize_image(image_path: str, max_width_inches: float = 4.5) -> str:
    """将图片等比缩放，返回处理后的临时路径。"""
    img = Image.open(image_path)
    dpi = img.info.get("dpi", (96, 96))[0]
    max_px = int(max_width_inches * dpi)
    if img.width > max_px:
        ratio = max_px / img.width
        new_size = (max_px, int(img.height * ratio))
        img = img.resize(new_size, Image.LANCZOS)

    out_path = image_path + "_resized.png"
    img.save(out_path, "PNG")
    return out_path


def _replace_paragraph_text(para, placeholder: str, replacement: str):
    """在段落中替换占位符，保留原有格式。"""
    if placeholder in para.text:
        for run in para.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)
                return
        # 若占位符跨 run，直接重写段落文本
        full_text = para.text.replace(placeholder, replacement)
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = full_text


def _insert_image_after_paragraph(doc: Document, para, image_path: str, width_inches: float = 4.5):
    """在指定段落后插入图片。"""
    resized = _resize_image(image_path, width_inches)
    new_para = OxmlElement("w:p")
    para._element.addnext(new_para)

    from docx.oxml.ns import nsmap
    run_elem = OxmlElement("w:r")
    new_para.append(run_elem)

    # 使用 python-docx 高层 API 在文档末尾添加图片段落，再移动到目标位置
    img_para = doc.add_paragraph()
    run = img_para.add_run()
    run.add_picture(resized, width=Inches(width_inches))
    img_para._element.getparent().remove(img_para._element)
    para._element.addnext(img_para._element)

    import os
    if os.path.exists(resized):
        os.remove(resized)


def render_report(
    output_path: str,
    payload,
    sds_page_path: str | None,
    mass_spec_path: str | None,
):
    """基于模板渲染完整的 Word 报告。"""
    _ensure_template()
    doc = Document(str(TEMPLATE_PATH))

    from datetime import datetime

    replacements = {
        "{{report_title}}": payload.report_title,
        "{{created_at}}": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "{{vector_name}}": payload.materials.vector_name,
        "{{host_cell}}": payload.materials.host_cell,
        "{{target_gene_size_kda}}": str(payload.materials.target_gene_size_kda),
        "{{tags}}": ", ".join(payload.materials.tags),
        "{{reagents_table}}": "; ".join(
            f"{r.get('name', '')}（{r.get('manufacturer', '')}，批号：{r.get('lot', '')}）"
            for r in payload.materials.reagents
        ),
        "{{instruments_table}}": "; ".join(
            f"{i.get('name', '')}（{i.get('manufacturer', '')}，编号：{i.get('serial', '')}）"
            for i in payload.materials.instruments
        ),
        "{{media_type}}": payload.fermentation.media_type,
        "{{inoculation_volume_percent}}": str(payload.fermentation.inoculation_volume_percent),
        "{{culture_temperature}}": str(payload.fermentation.culture_temperature),
        "{{shaker_speed_rpm}}": str(payload.fermentation.shaker_speed_rpm),
        "{{induction_od600}}": str(payload.fermentation.induction_od600),
        "{{iptg_concentration_mm}}": str(payload.fermentation.iptg_concentration_mm),
        "{{induction_temperature}}": str(payload.fermentation.induction_temperature),
        "{{expression_form}}": "可溶性表达" if payload.expression.is_soluble else "包涵体表达",
        "{{inclusion_body_treatment}}": payload.expression.inclusion_body_treatment or "N/A",
        "{{purification_method}}": payload.purification.purification_method,
        "{{assay_method}}": payload.purification.assay_method,
        "{{protein_concentration_mg_ml}}": str(payload.purification.protein_concentration_mg_ml),
        "{{protein_purity_percent}}": str(payload.purification.protein_purity_percent),
    }

    image_placeholders = {
        "{{sds_page_image}}": sds_page_path,
        "{{mass_spec_image}}": mass_spec_path,
    }

    for para in doc.paragraphs:
        for placeholder, value in replacements.items():
            _replace_paragraph_text(para, placeholder, value)

        for placeholder, img_path in image_placeholders.items():
            if placeholder in para.text:
                _replace_paragraph_text(para, placeholder, "")
                if img_path:
                    _insert_image_after_paragraph(doc, para, img_path)

    doc.save(output_path)
