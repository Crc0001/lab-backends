from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import os
import re
import tempfile
import base64
import html as html_module
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

UPLOAD_DIR = Path(os.environ.get("UPLOAD_DIR", "static/uploads"))
REPORTS_DIR = Path(os.environ.get("REPORTS_DIR", "static/reports"))

_SONG = "宋体"
_TNR = "Times New Roman"
_PT_BODY = Pt(12)
_PT_TABLE = Pt(10.5)
_PT_H1 = Pt(12)


def _plot_linearity(xs, ys, equation: str, r2: float) -> str:
    plt.rcParams['font.sans-serif'] = ['SimSun', 'Microsoft YaHei', 'DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False
    fig, ax = plt.subplots(figsize=(5, 3.5))
    ax.scatter(xs, ys, color='steelblue', zorder=3)
    m = np.polyfit(xs, ys, 1)
    x_line = np.linspace(min(xs), max(xs), 100)
    ax.plot(x_line, np.polyval(m, x_line), 'r--', linewidth=1)
    ax.set_xlabel('浓度 (μg/mL)', fontsize=9)
    ax.set_ylabel('OD值', fontsize=9)
    ax.grid(True, linestyle='--', alpha=0.4)
    # 公式标注在图内右下角
    ax.text(0.97, 0.05, f'{equation}\nR²={r2:.4f}',
            transform=ax.transAxes, fontsize=8,
            ha='right', va='bottom',
            bbox=dict(boxstyle='round,pad=0.3', facecolor='lightyellow', alpha=0.8))
    plt.tight_layout()
    tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
    fig.savefig(tmp.name, dpi=150)
    plt.close(fig)
    return tmp.name


def _set_run_font(run, bold=False, size=None):
    run.bold = bold
    if size:
        run.font.size = size
    run.font.name = _TNR
    run._element.rPr.rFonts.set(qn("w:eastAsia"), _SONG)


def _add_h1(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, bold=True, size=_PT_H1)


def _add_h2(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, bold=True, size=_PT_BODY)


def _set_cell_text(cell, text: str, bold=False):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    cell.paragraphs[0].clear()
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = cell.paragraphs[0].add_run(text)
    _set_run_font(run, bold=bold, size=_PT_TABLE)


def _add_table(doc, headers: list, rows: list):
    col_count = len(headers)
    if col_count == 0:
        return
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row[:col_count]):
            _set_cell_text(table.rows[r_idx + 1].cells[c_idx], val)
    doc.add_paragraph()


def _add_procedure_groups_docx(doc, titles, sub_titles_groups, header_groups, table_groups):
    default_proc_headers = ["序号", "步骤名称", "具体步骤"]
    for i, title in enumerate(titles):
        _add_h2(doc, title)
        sub_titles = sub_titles_groups[i] if i < len(sub_titles_groups) else []
        headers_group = header_groups[i] if i < len(header_groups) else []
        tables_group = table_groups[i] if i < len(table_groups) else []
        max_len = max(len(sub_titles), len(headers_group), len(tables_group), 1)
        for j in range(max_len):
            sub_title = sub_titles[j] if j < len(sub_titles) else f"小小标题{j + 1}"
            hdrs = headers_group[j] if j < len(headers_group) else default_proc_headers
            rows = tables_group[j] if j < len(tables_group) else []
            p = doc.add_paragraph()
            run = p.add_run(sub_title)
            _set_run_font(run, bold=True, size=_PT_TABLE)
            _add_table(doc, hdrs, rows)


def _procedure_groups_html(titles, sub_titles_groups, header_groups, table_groups):
    e = html_module.escape
    default_proc_headers = ["序号", "步骤名称", "具体步骤"]
    blocks = []
    for i, title in enumerate(titles):
        blocks.append(f"<h3>{e(title)}</h3>")
        sub_titles = sub_titles_groups[i] if i < len(sub_titles_groups) else []
        headers_group = header_groups[i] if i < len(header_groups) else []
        tables_group = table_groups[i] if i < len(table_groups) else []
        max_len = max(len(sub_titles), len(headers_group), len(tables_group), 1)
        for j in range(max_len):
            sub_title = sub_titles[j] if j < len(sub_titles) else f"小小标题{j + 1}"
            hdrs = headers_group[j] if j < len(headers_group) else default_proc_headers
            rows = tables_group[j] if j < len(tables_group) else []
            blocks.append(f"<h4>{e(sub_title)}</h4>")
            blocks.append(_table_html_simple(hdrs, rows))
    return blocks


def _table_html_simple(headers: list, rows: list) -> str:
    if not headers:
        return ""
    e = html_module.escape
    ths = "".join(f"<th>{e(h)}</th>" for h in headers)
    trs = ""
    for row in rows:
        tds = "".join(f"<td>{e(str(v))}</td>" for v in row)
        trs += f"<tr>{tds}</tr>"
    return f"<table><thead><tr>{ths}</tr></thead><tbody>{trs}</tbody></table>"


def _add_procedure_groups_pdf(story, paragraph_cls, tbl_func, h2_style, body_style, titles, sub_titles_groups, header_groups, table_groups):
    default_proc_headers = ["序号", "步骤名称", "具体步骤"]
    for i, title in enumerate(titles):
        story.append(paragraph_cls(title, h2_style))
        sub_titles = sub_titles_groups[i] if i < len(sub_titles_groups) else []
        headers_group = header_groups[i] if i < len(header_groups) else []
        tables_group = table_groups[i] if i < len(table_groups) else []
        max_len = max(len(sub_titles), len(headers_group), len(tables_group), 1)
        for j in range(max_len):
            sub_title = sub_titles[j] if j < len(sub_titles) else f"小小标题{j + 1}"
            story.append(paragraph_cls(sub_title, body_style))
            hdrs = headers_group[j] if j < len(headers_group) else default_proc_headers
            rows = tables_group[j] if j < len(tables_group) else []
            story.extend(tbl_func(hdrs, rows))


def _find_upload_file(file_id: str | None) -> Path | None:
    if not file_id:
        return None
    path = UPLOAD_DIR / file_id
    return path if path.exists() else None


def _approver_display(payload, idx: int, approver_text: str) -> str:
    signatures = getattr(payload, "approver_signatures", None) or []
    if idx < len(signatures) and signatures[idx]:
        return "[签名]"
    return approver_text


def _insert_signature_into_cell(cell, file_id: str | None, width_inches: float = 1.4):
    img_path = _find_upload_file(file_id)
    if not img_path:
        return False
    resized = _resize_image(str(img_path), width_inches)
    try:
        para = cell.paragraphs[0]
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(resized, width=Inches(width_inches))
        return True
    finally:
        if os.path.exists(resized):
            os.remove(resized)


def _signature_img_html(file_id: str | None, max_width: int = 140, max_height: int = 60) -> str:
    img_path = _find_upload_file(file_id)
    if not img_path:
        return ""
    try:
        img_b64 = base64.b64encode(img_path.read_bytes()).decode()
        suffix = img_path.suffix.lower().lstrip('.')
        mime = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png", "bmp": "bmp", "tiff": "tiff"}.get(suffix, "png")
        return f'<img src="data:image/{mime};base64,{img_b64}" style="max-width:{max_width}px;max-height:{max_height}px">'
    except Exception:
        return ""


def _resize_image(image_path: str, max_width_inches: float = 4.0) -> str:
    img = Image.open(image_path)
    dpi = img.info.get("dpi", (96, 96))[0]
    max_px = int(max_width_inches * dpi)
    if img.width > max_px:
        ratio = max_px / img.width
        img = img.resize((max_px, int(img.height * ratio)), Image.LANCZOS)
    out = image_path + "_resized.png"
    img.save(out, "PNG")
    return out


def _clean_report_text(value: str | None) -> str:
    text = (value or "").replace("\r\n", "\n").replace("\r", "\n")
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)


def _split_text_chunks(value: str, chunk_size: int) -> list[str]:
    text = _clean_report_text(value)
    if not text:
        return [""]
    return [text[i:i + chunk_size] for i in range(0, min(len(text), 12000), chunk_size)]


def _find_appendix_txt_file(fpath: str | None) -> Path | None:
    if not fpath:
        return None
    candidates: list[Path] = []
    raw = Path(fpath)
    raw_str = str(fpath).strip()
    candidates.append(raw)
    candidates.append(UPLOAD_DIR / raw.name)
    if raw.suffix.lower() != ".txt":
        candidates.append(UPLOAD_DIR / f"{raw.name}.txt")
    for match in re.findall(r"[0-9a-fA-F-]{32,}\.txt|[0-9a-fA-F-]{32,}", raw_str):
        candidates.append(UPLOAD_DIR / match)
        if not match.lower().endswith('.txt'):
            candidates.append(UPLOAD_DIR / f"{match}.txt")
    seen = set()
    for candidate in candidates:
        key = str(candidate)
        if key in seen:
            continue
        seen.add(key)
        if candidate.exists() and candidate.suffix.lower() == ".txt":
            return candidate
    return None


def _append_txt_attachments_docx(doc, file_paths: list[str] | None):
    for fpath in (file_paths or []):
        p_path = _find_appendix_txt_file(fpath)
        if p_path is None:
            continue
        _add_h2(doc, f"附件：{p_path.name}")
        try:
            content = p_path.read_text(encoding='utf-8', errors='replace')
            for chunk in _split_text_chunks(content, 2000):
                p = doc.add_paragraph()
                run = p.add_run(chunk)
                _set_run_font(run, size=Pt(10))
        except Exception:
            continue
        doc.add_paragraph()


def _append_txt_attachments_html(sections: list[str], file_paths: list[str] | None, escape_fn):
    for fpath in (file_paths or []):
        p_path = _find_appendix_txt_file(fpath)
        if p_path is None:
            continue
        try:
            content = p_path.read_text(encoding='utf-8', errors='replace')
            sections.append(f"<h3>附件：{escape_fn(p_path.name)}</h3>")
            sections.append(f"<pre style='white-space:pre-wrap;font-size:12px;background:#f9f9f9;padding:8px;border:1px solid #ddd'>{escape_fn(_clean_report_text(content)[:12000])}</pre>")
        except Exception:
            continue


def _append_txt_attachments_pdf(story, paragraph_cls, body_style, file_paths: list[str] | None):
    import xml.sax.saxutils as _sax
    for fpath in (file_paths or []):
        p_path = _find_appendix_txt_file(fpath)
        if p_path is None:
            continue
        try:
            content = p_path.read_text(encoding='utf-8', errors='replace')
            story.append(paragraph_cls(f"附件：{_sax.escape(p_path.name)}", body_style))
            for chunk in _split_text_chunks(content, 1800):
                story.append(paragraph_cls(_sax.escape(chunk).replace('\n', '<br/>'), body_style))
            story.append(Spacer(1, 6))
        except Exception:
            continue


def _add_page_numbers(doc, first_page_footer: str = ""):
    """为 Word 文档添加页码：首页页脚用自定义文字，其余页显示"第 X 页"。"""
    def _make_fld(fld_type: str):
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), fld_type)
        return fld

    def _make_instrText(text: str):
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = text
        return instr

    def _page_num_run(para):
        """在段落中插入 PAGE 域（当前页）。"""
        r = OxmlElement('w:r')
        r.append(_make_fld('begin'))
        r.append(_make_instrText(' PAGE '))
        r.append(_make_fld('separate'))
        r.append(_make_fld('end'))
        para._p.append(r)

    section = doc.sections[0]
    section.different_first_page_header_footer = True

    first_footer = section.first_page_footer
    fp = first_footer.paragraphs[0] if first_footer.paragraphs else first_footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(first_page_footer or "")
    _set_run_font(run, size=Pt(10))

    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pre = p.add_run("第 ")
    _set_run_font(run_pre, size=Pt(10))
    _page_num_run(p)
    run_suf = p.add_run(" 页")
    _set_run_font(run_suf, size=Pt(10))


def _add_table(doc, headers: list, rows: list):
    col_count = len(headers)
    if col_count == 0:
        return
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row[:col_count]):
            _set_cell_text(table.rows[r_idx + 1].cells[c_idx], val)
    doc.add_paragraph()


def _render_qc_section(doc, payload):
    """渲染 QC 数据处理 + 结论（供正文和附录复用），按实际内容输出。"""
    if payload.qc_concentrations or payload.qc_od_values:
        _add_h2(doc, "附表1 数据处理")
        headers = payload.qc_data_headers or []
        conc_row = ["浓度(mg/ml)"] + list(payload.qc_concentrations or [])
        od_row = ["OD545nm"] + list(payload.qc_od_values or [])
        if headers:
            _add_table(doc, headers, [conc_row, od_row])
        if payload.qc_linear_equation:
            p = doc.add_paragraph()
            run = p.add_run(f"线性方程：{payload.qc_linear_equation}")
            _set_run_font(run, size=_PT_BODY)
        if payload.qc_correlation_coeff is not None:
            p = doc.add_paragraph()
            run = p.add_run(f"相关系数 R：{payload.qc_correlation_coeff:.6f}")
            _set_run_font(run, size=_PT_BODY)

    if payload.qc_conclusion_headers and payload.qc_conclusion_rows:
        _add_h2(doc, "附表2 结论")
        try:
            xs = [float(v) for v in (payload.qc_concentrations or []) if v]
            ys = [float(v) for v in (payload.qc_od_values or []) if v]
            if len(xs) >= 2 and len(xs) == len(ys):
                chart_path = _plot_linearity(xs, ys,
                    payload.qc_linear_equation or '',
                    payload.qc_correlation_coeff or 0.0)
                doc.add_picture(chart_path, width=Inches(4.5))
                os.remove(chart_path)
        except Exception:
            pass
        row_labels = payload.qc_conclusion_row_labels or ['OD545nm', '浓度(mg/ml)']
        headers = payload.qc_conclusion_headers
        col_count = len(headers)
        n_data = col_count - 1
        od_vals = payload.qc_conclusion_rows[0] if len(payload.qc_conclusion_rows) > 0 else []
        conc_vals = payload.qc_conclusion_rows[1] if len(payload.qc_conclusion_rows) > 1 else []
        data_rows = [
            [row_labels[0] if row_labels else 'OD545nm'] + list(od_vals[:n_data]),
            [row_labels[1] if len(row_labels) > 1 else '浓度(mg/ml)'] + list(conc_vals[:n_data]),
        ]
        _add_table(doc, headers, data_rows)


def render_report_v2(output_path: str, payload) -> None:
    doc = Document()

    # ── 封面 ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(payload.report_title or "蛋白发酵实验报告")
    _set_run_font(run, bold=True, size=Pt(22))

    for label, val in [
        ("文档编号", payload.doc_number),
        ("密级", payload.security_level),
        ("项目名称", payload.project_name),
        ("生成日期", datetime.now().strftime("%Y-%m-%d")),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}：{val}")
        _set_run_font(run, size=_PT_BODY)
    doc.add_paragraph()

    _add_h2(doc, "实验人员信息")
    personnel_rows = [[p.editor, p.reviewer, _approver_display(payload, idx, p.approver), p.date, p.version]
                     for idx, p in enumerate(payload.personnel)]
    _add_table(doc, ["编制人", "审核人", "批准人", "编制日期", "版本号"], personnel_rows)
    for idx, signature in enumerate(getattr(payload, "approver_signatures", None) or []):
        if idx < len(payload.personnel):
            _insert_signature_into_cell(doc.tables[-1].rows[idx + 1].cells[2], signature)

    # ── 一、材料试剂与仪器 ──
    _add_h1(doc, "一、材料试剂与仪器")
    _add_h2(doc, "1. 菌株与质粒")
    _add_table(doc, payload.strains_headers or ["名称与信息", "厂家", "厂家生产日期", "有效期"], payload.strains)
    _add_h2(doc, "2. 试剂信息")
    _add_table(doc, payload.reagents_headers or ["名称", "厂家", "批号"], payload.reagents)
    _add_h2(doc, "3. 仪器信息")
    _add_table(doc, payload.instruments_headers or ["名称", "厂家", "批号", "校准日期"], payload.instruments)
    _add_h2(doc, "4. 现用现配")
    _add_table(doc, payload.reagent_config_headers or ["名称", "配置日期", "配方", "有效期"], payload.reagent_configs)
    _add_h2(doc, "5. 耗材清单")
    _add_table(doc, payload.consumables_headers or ["名称", "厂家", "生产日期", "有效期"], payload.consumables)

    # ── 二、操作步骤 ──
    _add_h1(doc, "二、操作步骤")
    dyn_titles = payload.dynamic_section_titles or []
    _add_procedure_groups_docx(
        doc,
        [f"{i+1}. {t}" for i, t in enumerate(dyn_titles)],
        payload.dynamic_section_sub_titles or [],
        payload.dynamic_section_headers or [],
        payload.dynamic_section_tables or [],
    )

    # ── 三、实验结果与分析 ──
    _add_h1(doc, "三、实验结果与分析")
    result_headers = ["序号", "步骤名称", "图谱"] + list(payload.result_extra_headers or [])
    result_table = doc.add_table(rows=1 + len(payload.results), cols=len(result_headers))
    result_table.style = "Table Grid"
    for i, h in enumerate(result_headers):
        _set_cell_text(result_table.rows[0].cells[i], h, bold=True)
    for r_idx, row in enumerate(payload.results):
        _set_cell_text(result_table.rows[r_idx + 1].cells[0], row.seq)
        _set_cell_text(result_table.rows[r_idx + 1].cells[1], row.name)
        for extra_idx, val in enumerate(row.extra_cols or []):
            cell_idx = 3 + extra_idx
            if cell_idx < len(result_headers):
                _set_cell_text(result_table.rows[r_idx + 1].cells[cell_idx], val)
        if row.image_file_id:
            img_path = UPLOAD_DIR / row.image_file_id
            if img_path.exists():
                resized = _resize_image(str(img_path), 2.5)
                result_table.rows[r_idx + 1].cells[2].paragraphs[0].add_run().add_picture(resized, width=Inches(2.5))
                if os.path.exists(resized):
                    os.remove(resized)
    doc.add_paragraph()

    # ── 四、实验结论与工艺建议 ──
    _add_h1(doc, "四、实验结论与工艺建议")
    if payload.conclusion_items:
        for i, item in enumerate(payload.conclusion_items):
            p = doc.add_paragraph()
            run = p.add_run(f"({i + 1}) {item}")
            _set_run_font(run, size=_PT_BODY)
    else:
        conclusion_headers = payload.conclusion_table_headers or ["实验结论", "工艺建议"]
        _add_table(doc, conclusion_headers, payload.conclusion_table_rows or [])

    # ── 附录（新页）──
    doc.add_page_break()
    _add_h1(doc, "附录")
    _render_qc_section(doc, payload)

    # 附件 TXT
    _append_txt_attachments_docx(doc, payload.appendix_file_paths)

    # ── 页码 ──
    _add_page_numbers(doc, payload.first_page_footer)

    doc.save(output_path)


def render_preview_html(payload) -> str:
    """生成 HTML 预览字符串。"""
    e = html_module.escape

    def table_html(headers: list, rows: list) -> str:
        if not headers:
            return ""
        ths = "".join(f"<th>{e(h)}</th>" for h in headers)
        trs = ""
        for row in rows:
            tds = "".join(f"<td>{v if isinstance(v, str) and '<img' in v else e(str(v))}</td>" for v in row)
            trs += f"<tr>{tds}</tr>"
        return f"<table><thead><tr>{ths}</tr></thead><tbody>{trs}</tbody></table>"

    sections = []

    # 封面
    sections.append(f"<h1 class='cover'>{e(payload.report_title or '蛋白发酵实验报告')}</h1>")
    sections.append(f"<p>文档编号：{e(payload.doc_number)} &nbsp; 密级：{e(payload.security_level)}</p>")
    sections.append(f"<p>项目名称：{e(payload.project_name)} &nbsp; 生成日期：{datetime.now().strftime('%Y-%m-%d')}</p>")

    # 人员
    sections.append("<h2>实验人员信息</h2>")
    personnel_rows = []
    for idx, p in enumerate(payload.personnel):
        signature_html = _signature_img_html((getattr(payload, "approver_signatures", None) or [None] * len(payload.personnel))[idx])
        personnel_rows.append([p.editor, p.reviewer, signature_html or p.approver, p.date, p.version])
    sections.append(table_html(
        ["编制人", "审核人", "批准人", "编制日期", "版本号"],
        personnel_rows
    ))

    # 一
    sections.append("<h2>一、材料试剂与仪器</h2>")
    sections.append("<h3>1. 菌株与质粒</h3>")
    sections.append(table_html(payload.strains_headers or ["名称与信息", "厂家", "厂家生产日期", "有效期"], payload.strains))
    sections.append("<h3>2. 试剂信息</h3>")
    sections.append(table_html(payload.reagents_headers or ["名称", "厂家", "批号"], payload.reagents))
    sections.append("<h3>3. 仪器信息</h3>")
    sections.append(table_html(payload.instruments_headers or ["名称", "厂家", "批号", "校准日期"], payload.instruments))
    sections.append("<h3>4. 现用现配</h3>")
    sections.append(table_html(payload.reagent_config_headers or ["名称", "配置日期", "配方", "有效期"], payload.reagent_configs))
    sections.append("<h3>5. 耗材清单</h3>")
    sections.append(table_html(payload.consumables_headers or ["名称", "厂家", "生产日期", "有效期"], payload.consumables))

    # 二
    sections.append("<h2>二、操作步骤</h2>")
    dyn_titles = payload.dynamic_section_titles or []
    sections.extend(_procedure_groups_html(
        [f"{i+1}. {t}" for i, t in enumerate(dyn_titles)],
        payload.dynamic_section_sub_titles or [],
        payload.dynamic_section_headers or [],
        payload.dynamic_section_tables or [],
    ))

    # 三
    sections.append("<h2>三、实验结果与分析</h2>")
    result_headers = ["序号", "步骤名称", "图谱"] + list(payload.result_extra_headers or [])
    ths = "".join(f"<th>{e(h)}</th>" for h in result_headers)
    result_trs = ""
    for r in payload.results:
        img_cell = ""
        if r.image_file_id:
            img_path = UPLOAD_DIR / r.image_file_id
            if img_path.exists():
                try:
                    img_b64 = base64.b64encode(img_path.read_bytes()).decode()
                    suffix = img_path.suffix.lower().lstrip(".")
                    mime = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png", "bmp": "bmp", "tiff": "tiff"}.get(suffix, "png")
                    img_cell = f'<img src="data:image/{mime};base64,{img_b64}" style="max-width:200px;max-height:160px">'
                except Exception:
                    img_cell = "(图谱)"
            else:
                img_cell = "(图谱)"
        extra_tds = "".join(f"<td>{e(str(v))}</td>" for v in (r.extra_cols or []))
        result_trs += f"<tr><td>{e(r.seq)}</td><td>{e(r.name)}</td><td>{img_cell}</td>{extra_tds}</tr>"
    sections.append(f"<table><thead><tr>{ths}</tr></thead><tbody>{result_trs}</tbody></table>")

    # 四
    sections.append("<h2>四、实验结论与工艺建议</h2>")
    sections.append(table_html(
        payload.conclusion_table_headers or ["实验结论", "工艺建议"],
        payload.conclusion_table_rows or []
    ))

    # 附录
    sections.append("<hr><h2>附录</h2>")
    if payload.qc_concentrations or payload.qc_od_values:
        sections.append("<h3>数据处理</h3>")
        headers = payload.qc_data_headers or []
        if headers:
            sections.append(table_html(headers, [
                ["浓度(mg/ml)"] + list(payload.qc_concentrations or []),
                ["OD545nm"] + list(payload.qc_od_values or []),
            ]))
        if payload.qc_linear_equation:
            sections.append(f"<p>线性方程：{e(payload.qc_linear_equation)}</p>")
        if payload.qc_correlation_coeff is not None:
            sections.append(f"<p>相关系数 R：{payload.qc_correlation_coeff:.6f}</p>")
    if payload.qc_conclusion_headers and payload.qc_conclusion_rows:
        sections.append("<h3>结论</h3>")
        # 线性图：生成后转 base64 嵌入
        try:
            xs = [float(v) for v in (payload.qc_concentrations or []) if v]
            ys = [float(v) for v in (payload.qc_od_values or []) if v]
            if len(xs) >= 2 and len(xs) == len(ys):
                chart_path = _plot_linearity(xs, ys,
                    payload.qc_linear_equation or '',
                    payload.qc_correlation_coeff or 0.0)
                img_b64 = base64.b64encode(Path(chart_path).read_bytes()).decode()
                os.remove(chart_path)
                sections.append(f'<img src="data:image/png;base64,{img_b64}" style="max-width:400px;display:block;margin:8px 0">')
        except Exception:
            pass
        row_labels = payload.qc_conclusion_row_labels or ['OD545nm', '浓度(mg/ml)']
        col_count = len(payload.qc_conclusion_headers)
        n_data = col_count - 1
        od_vals = payload.qc_conclusion_rows[0] if payload.qc_conclusion_rows else []
        conc_vals = payload.qc_conclusion_rows[1] if len(payload.qc_conclusion_rows) > 1 else []
        sections.append(table_html(payload.qc_conclusion_headers, [
            [row_labels[0]] + list(od_vals[:n_data]),
            [row_labels[1] if len(row_labels) > 1 else '浓度(mg/ml)'] + list(conc_vals[:n_data]),
        ]))

    # 附录 TXT
    _append_txt_attachments_html(sections, payload.appendix_file_paths, e)

    body = "\n".join(sections)
    return f"""<!DOCTYPE html>
<html lang="zh"><head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{e(payload.report_title or '报告预览')}</title>
<style>
body{{font-family:'宋体',serif;font-size:14px;margin:16px;line-height:1.6;color:#222}}
h1.cover{{text-align:center;font-size:22px}}
h2{{font-size:15px;border-bottom:1px solid #ccc;padding-bottom:2px;margin-top:20px}}
h3{{font-size:14px;margin-top:12px}}
table{{border-collapse:collapse;width:100%;margin-bottom:12px;font-size:13px}}
th,td{{border:1px solid #bbb;padding:4px 8px;text-align:left}}
th{{background:#f5f5f5;font-weight:bold}}
hr{{margin:24px 0}}
p{{margin:4px 0}}
</style>
</head><body>
{body}
</body></html>"""


def render_pdf(output_path: str, payload) -> None:
    """用 reportlab 直接生成 PDF，不依赖 LibreOffice。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, PageBreak, Image as RLImage)
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import glob as _glob

    # 注册中文字体，按平台搜索
    font_name = "Helvetica"
    font_candidates = [
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "/system/fonts/NotoSansCJK-Regular.ttc",
        "/system/fonts/DroidSansFallback.ttf",
        "/system/fonts/NotoSerifCJK-Regular.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttf",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttf",
        "/System/Library/Fonts/PingFang.ttc",
    ]
    # 追加环境变量指定的字体路径
    env_font = os.environ.get("CJK_FONT_PATH")
    if env_font:
        font_candidates.insert(0, env_font)
    for pattern in font_candidates:
        found = _glob.glob(pattern)
        if found:
            try:
                pdfmetrics.registerFont(TTFont("CJK", found[0]))
                font_name = "CJK"
                break
            except Exception:
                continue

    h1 = ParagraphStyle("h1", fontName=font_name, fontSize=13, spaceAfter=6, spaceBefore=12, leading=18, textColor=colors.HexColor("#1a3a6b"))
    h2 = ParagraphStyle("h2", fontName=font_name, fontSize=11, spaceAfter=4, spaceBefore=8, leading=16)
    body = ParagraphStyle("body", fontName=font_name, fontSize=10, spaceAfter=3, leading=14)
    cover = ParagraphStyle("cover", fontName=font_name, fontSize=18, alignment=1, spaceAfter=8, leading=24)

    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2.5*cm, bottomMargin=2.5*cm)

    import xml.sax.saxutils as _sax
    cell_style = ParagraphStyle("cell", fontName=font_name, fontSize=9, leading=12, wordWrap='CJK')

    def _cell(v):
        if isinstance(v, str):
            return Paragraph(_sax.escape(v), cell_style)
        return v

    def tbl(headers, rows):
        if not headers:
            return []
        data = [[_cell(h) for h in headers]] + (
            [[_cell(v) for v in row] for row in rows] if rows else [[_cell("") for _ in headers]]
        )
        col_max = [max(len(str(h or '')), *(len(str(row[i] if i < len(row) else '')) for row in rows), 4) for i, h in enumerate(headers)]
        weights = [min(max(length, 4), 30) for length in col_max]
        total_weight = sum(weights) or len(headers)
        col_widths = [doc.width * weight / total_weight for weight in weights]
        t = Table(data, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#e8edf5")),
            ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("WORDWRAP", (0,0), (-1,-1), "CJK"),
        ]))
        return [t, Spacer(1, 6)]

    story = []
    temp_signature_images = []
    temp_images = []

    # 封面
    story.append(Paragraph(_sax.escape(payload.report_title or "蛋白发酵实验报告"), cover))
    for label, val in [("文档编号", payload.doc_number), ("密级", payload.security_level),
                       ("项目名称", payload.project_name),
                       ("生成日期", datetime.now().strftime("%Y-%m-%d"))]:
        story.append(Paragraph(f"{_sax.escape(label)}：{_sax.escape(str(val or ''))}", body))
    story.append(Spacer(1, 12))

    # 人员
    story.append(Paragraph("实验人员信息", h2))
    personnel_data = [[_cell("编制人"), _cell("审核人"), _cell("批准人"), _cell("编制日期"), _cell("版本号")]]
    signatures = getattr(payload, "approver_signatures", None) or []
    for idx, p in enumerate(payload.personnel):
        signature_value = _cell(p.approver)
        if idx < len(signatures):
            img_path = _find_upload_file(signatures[idx])
            if img_path:
                resized = _resize_image(str(img_path), 1.4)
                temp_signature_images.append(resized)
                signature_value = RLImage(resized, width=72, height=32)
        personnel_data.append([_cell(p.editor), _cell(p.reviewer), signature_value, _cell(p.date), _cell(p.version)])
    personnel_col_w = doc.width / 5
    personnel_tbl = Table(personnel_data, colWidths=[personnel_col_w] * 5, repeatRows=1)
    personnel_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#e8edf5")),
        ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("ALIGN", (0,0), (-1,-1), "CENTER"),
    ]))
    story.extend([personnel_tbl, Spacer(1, 6)])

    # 一
    story.append(Paragraph("一、材料试剂与仪器", h1))
    for title, hdrs, rows in [
        ("1. 菌株与质粒", payload.strains_headers or ["名称与信息","厂家","厂家生产日期","有效期"], payload.strains),
        ("2. 试剂信息", payload.reagents_headers or ["名称","厂家","批号"], payload.reagents),
        ("3. 仪器信息", payload.instruments_headers or ["名称","厂家","批号","校准日期"], payload.instruments),
        ("4. 现用现配", payload.reagent_config_headers or ["名称","配置日期","配方","有效期"], payload.reagent_configs),
        ("5. 耗材清单", payload.consumables_headers or ["名称","厂家","生产日期","有效期"], payload.consumables),
    ]:
        story.append(Paragraph(title, h2))
        story.extend(tbl(hdrs, rows))

    # 二
    story.append(Paragraph("二、操作步骤", h1))
    dyn_titles = payload.dynamic_section_titles or []
    _add_procedure_groups_pdf(
        story, Paragraph, tbl, h2, body,
        [f"{i+1}. {t}" for i, t in enumerate(dyn_titles)],
        payload.dynamic_section_sub_titles or [],
        payload.dynamic_section_headers or [],
        payload.dynamic_section_tables or [],
    )

    # 三
    story.append(Paragraph("三、实验结果与分析", h1))
    result_hdrs = [_cell("序号"), _cell("步骤名称"), _cell("图谱")] + [_cell(h) for h in (payload.result_extra_headers or [])]
    result_data = [result_hdrs]
    for r in payload.results:
        img_val = _cell("")
        if r.image_file_id:
            img_path = UPLOAD_DIR / r.image_file_id
            if img_path.exists():
                resized = _resize_image(str(img_path), 2.0)
                temp_images.append(resized)
                img_val = RLImage(resized, width=96, height=72)
            else:
                img_val = _cell("(图谱缺失)")
        row_vals = [_cell(r.seq), _cell(r.name), img_val] + [_cell(v) for v in (r.extra_cols or [])]
        result_data.append(row_vals)
    col_w = doc.width / max(len(result_hdrs), 1)
    result_tbl = Table(result_data, colWidths=[col_w] * len(result_hdrs), repeatRows=1)
    result_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#e8edf5")),
        ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("ALIGN", (0,0), (-1,-1), "CENTER"),
    ]))
    story.extend([result_tbl, Spacer(1, 6)])

    # 四
    story.append(Paragraph("四、实验结论与工艺建议", h1))
    story.extend(tbl(
        payload.conclusion_table_headers or ["实验结论","工艺建议"],
        payload.conclusion_table_rows or []
    ))

    # 附录
    story.append(PageBreak())
    story.append(Paragraph("附录", h1))
    if payload.qc_concentrations or payload.qc_od_values:
        story.append(Paragraph("数据处理", h2))
        hdrs = payload.qc_data_headers or []
        if hdrs:
            story.extend(tbl(hdrs, [
                ["浓度(mg/ml)"] + list(payload.qc_concentrations or []),
                ["OD545nm"] + list(payload.qc_od_values or []),
            ]))
        if payload.qc_linear_equation:
            story.append(Paragraph(f"线性方程：{_sax.escape(payload.qc_linear_equation)}", body))
        if payload.qc_correlation_coeff is not None:
            story.append(Paragraph(f"相关系数 R：{payload.qc_correlation_coeff:.6f}", body))
    if payload.qc_conclusion_headers and payload.qc_conclusion_rows:
        story.append(Paragraph("结论", h2))
        try:
            xs = [float(v) for v in (payload.qc_concentrations or []) if v not in (None, "")]
            ys = [float(v) for v in (payload.qc_od_values or []) if v not in (None, "")]
            if len(xs) >= 2 and len(xs) == len(ys):
                chart_path = _plot_linearity(xs, ys,
                    payload.qc_linear_equation or '',
                    payload.qc_correlation_coeff or 0.0)
                chart_img = RLImage(chart_path, width=320, height=224)
                story.extend([chart_img, Spacer(1, 6)])
                temp_images.append(chart_path)
        except Exception:
            pass
        rl = payload.qc_conclusion_row_labels or ["OD545nm","浓度(mg/ml)"]
        n_data = len(payload.qc_conclusion_headers) - 1
        od_v = payload.qc_conclusion_rows[0] if payload.qc_conclusion_rows else []
        co_v = payload.qc_conclusion_rows[1] if len(payload.qc_conclusion_rows) > 1 else []
        story.extend(tbl(payload.qc_conclusion_headers, [
            [rl[0]] + list(od_v[:n_data]),
            [rl[1] if len(rl)>1 else "浓度(mg/ml)"] + list(co_v[:n_data]),
        ]))

    _append_txt_attachments_pdf(story, Paragraph, body, payload.appendix_file_paths)

    # 首页页脚页码（通过 onFirstPage / onLaterPages 回调）
    footer_text = payload.first_page_footer or ""

    def on_first_page(canvas, _):
        canvas.saveState()
        canvas.setFont(font_name, 9)
        canvas.drawCentredString(A4[0]/2, 1.2*cm, footer_text)
        canvas.restoreState()

    def on_later_pages(canvas, _):
        canvas.saveState()
        canvas.setFont(font_name, 9)
        canvas.drawCentredString(A4[0]/2, 1.2*cm, f"第 {canvas.getPageNumber()} 页")
        canvas.restoreState()

    try:
        doc.build(story, onFirstPage=on_first_page, onLaterPages=on_later_pages)
    finally:
        for temp_path in [*temp_signature_images, *temp_images]:
            if os.path.exists(temp_path):
                os.remove(temp_path)
