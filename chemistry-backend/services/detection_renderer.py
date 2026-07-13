import os
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn


def _make_linearity_chart_png(s43_cols) -> bytes | None:
    """用 matplotlib 生成散点+回归线图，返回 PNG bytes；数据不足时返回 None。"""
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import matplotlib.font_manager as fm
        import numpy as np

        # 尝试加载系统中文字体
        _cjk_fonts = [
            "C:/Windows/Fonts/simhei.ttf",
            "C:/Windows/Fonts/simsun.ttc",
            "C:/Windows/Fonts/msyh.ttc",
        ]
        _font_prop = None
        for _fp in _cjk_fonts:
            if os.path.exists(_fp):
                _font_prop = fm.FontProperties(fname=_fp, size=9)
                plt.rcParams['font.family'] = fm.FontProperties(fname=_fp).get_name()
                break

        pts = []
        for col in s43_cols:
            try:
                x = float(col[0])
                y = float(col[1])
                pts.append((x, y))
            except (ValueError, IndexError):
                pass
        if len(pts) < 2:
            return None

        pts.sort()
        xs = np.array([p[0] for p in pts])
        ys = np.array([p[1] for p in pts])
        coeffs = np.polyfit(xs, ys, 1)
        slope, intercept = coeffs[0], coeffs[1]
        x_line = np.linspace(xs.min(), xs.max(), 200)
        y_line = slope * x_line + intercept

        eq = s43_cols[0][2] if s43_cols and len(s43_cols[0]) > 2 else ''
        r_val = s43_cols[0][3] if s43_cols and len(s43_cols[0]) > 3 else ''

        fig, ax = plt.subplots(figsize=(5, 3.5))
        ax.scatter(xs, ys, color='royalblue', zorder=5,
                   label='数据点' if _font_prop else 'Data')
        ax.plot(x_line, y_line, color='darkorange', linewidth=1.5,
                label='回归线' if _font_prop else 'Fit')
        if eq or r_val:
            ax.set_title(f'{eq}  r={r_val}',
                         fontproperties=_font_prop, fontsize=9)
        ax.set_xlabel('浓度 (ng/ml)' if _font_prop else 'Concentration (ng/ml)',
                      fontproperties=_font_prop, fontsize=9)
        ax.set_ylabel('峰面积' if _font_prop else 'Peak Area',
                      fontproperties=_font_prop, fontsize=9)
        if _font_prop:
            ax.legend(prop=_font_prop, fontsize=8)
        else:
            ax.legend(fontsize=8)
        ax.grid(True, linestyle='--', alpha=0.5)
        plt.tight_layout()

        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=120)
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception:
        return None


_SONG = "宋体"
_TNR = "Times New Roman"
_PT_TITLE = Pt(22)   # 二号
_PT_H1 = Pt(16)      # 三号
_PT_BODY = Pt(12)    # 小四
_PT_TABLE = Pt(10.5) # 五号

_MS_LABELS = [
    '扫描模式', '离子源类型', '喷雾电压', '扫描类型', '一级质谱分辨率',
    '扫描范围', '时间范围', 'Sheath gas', 'Aug gas',
    'Capillary temp(℃)', 'Aux gas temp(℃)', '提取离子',
]


def _set_run(run, bold=False, size=None, color=None):
    run.bold = bold
    if size:
        run.font.size = size
    run.font.name = _TNR
    run._element.rPr.rFonts.set(qn("w:eastAsia"), _SONG)
    if color:
        run.font.color.rgb = color


def _para(doc, text, bold=False, size=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    _set_run(run, bold=bold, size=size or _PT_BODY)
    return p


def _h1(doc, text):
    _para(doc, text, bold=True, size=_PT_H1)


def _h2(doc, text):
    _para(doc, text, bold=True, size=_PT_BODY)


def _cell(cell, text, bold=False, center=True):
    para = cell.paragraphs[0]
    para.clear()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(str(text))
    _set_run(run, bold=bold, size=_PT_TABLE)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _table(doc, headers, rows, bold_header=True):
    if not headers:
        return
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        _cell(t.rows[0].cells[i], h, bold=bold_header)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row[:len(headers)]):
            _cell(t.rows[ri + 1].cells[ci], val)
    doc.add_paragraph()


def _s26_s27_table(doc, headers, rows, per_compound):
    """
    渲染 2.6/2.7 表格：按 per_compound 行一组，
    每组的"限度储备液浓度"列（col 1）纵向合并。
    """
    if not headers or not rows:
        _table(doc, headers, rows)
        return
    n_rows = len(rows)
    t = doc.add_table(rows=1 + n_rows, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        _cell(t.rows[0].cells[i], h, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row[:len(headers)]):
            _cell(t.rows[ri + 1].cells[ci], val)
    # 按组合并第1列（限度储备液浓度）
    group_start = 0
    while group_start < n_rows:
        group_end = min(group_start + per_compound - 1, n_rows - 1)
        if group_end > group_start:
            merge_cell = t.rows[group_start + 1].cells[1]
            for k in range(group_start + 1, group_end + 1):
                merge_cell = merge_cell.merge(t.rows[k + 1].cells[1])
        group_start += per_compound
    doc.add_paragraph()


def _s26_s27_split(doc, headers, rows, per_compound):
    """每 per_compound 行为一组，每组单独渲染一张表（独立表头）。"""
    if not rows:
        _table(doc, headers, [])
        return
    n_compounds = max(1, (len(rows) + per_compound - 1) // per_compound)
    for ci in range(n_compounds):
        chunk = rows[ci * per_compound: ci * per_compound + per_compound]
        if not chunk:
            continue
        # 每组恢复限度储备液浓度（第1列只有首行有值）
        limit_conc = chunk[0][1] if len(chunk[0]) > 1 else ''
        t = doc.add_table(rows=1 + len(chunk), cols=len(headers))
        t.style = "Table Grid"
        for i, h in enumerate(headers):
            _cell(t.rows[0].cells[i], h, bold=True)
        for ri, row in enumerate(chunk):
            for ci2, val in enumerate(row[:len(headers)]):
                _cell(t.rows[ri + 1].cells[ci2], val)
        # 纵向合并限度储备液浓度列（col 1）
        if len(chunk) > 1:
            mc = t.rows[1].cells[1]
            for k in range(1, len(chunk)):
                mc = mc.merge(t.rows[k + 1].cells[1])
            _cell(mc, limit_conc)
        doc.add_paragraph()


def _s41_calc_avg_rsd(rows):
    try:
        vals = [float(r[4]) for r in rows if len(r) > 4 and r[4] != '']
    except (ValueError, TypeError):
        vals = []
    if not vals:
        return '', ''
    avg = sum(vals) / len(vals)
    std = (sum((x - avg) ** 2 for x in vals) / max(len(vals) - 1, 1)) ** 0.5
    return f'{avg:.2f}', (f'{std / avg * 100:.2f}' if avg != 0 else '0.00')


def _s41_groups_table(doc, groups):
    headers = ['样本名称', '加入量\n(ng/ml)', '样本含量\n(ng/ml)', '测得量\n(ng/ml)', '回收率(%)', '平均回收率(%)', 'RSD(%)']
    if not groups:
        _table(doc, headers, [])
        return
    for g in groups:
        title = g.get('title', '')
        rows = g.get('rows', [])
        if title:
            _para(doc, title)
        if not rows:
            continue
        avg_str, rsd_str = _s41_calc_avg_rsd(rows)
        n_rows = len(rows)
        t = doc.add_table(rows=1 + n_rows, cols=7)
        t.style = "Table Grid"
        for i, h in enumerate(headers):
            _cell(t.rows[0].cells[i], h, bold=True)
        for ri, row in enumerate(rows):
            for ci in range(min(5, len(row))):
                _cell(t.rows[ri + 1].cells[ci], row[ci])
        if n_rows > 1:
            mc5 = t.rows[1].cells[5]
            mc6 = t.rows[1].cells[6]
            for k in range(1, n_rows):
                mc5 = mc5.merge(t.rows[k + 1].cells[5])
                mc6 = mc6.merge(t.rows[k + 1].cells[6])
            _cell(mc5, avg_str)
            _cell(mc6, rsd_str)
        else:
            _cell(t.rows[1].cells[5], avg_str)
            _cell(t.rows[1].cells[6], rsd_str)
        doc.add_paragraph()


def _s43_groups_from_payload(p):
    """从 payload 提取 s43_groups 列表，每项为 {title, headers, cols}。
    优先读新格式 s43_groups，兼容旧格式 s43_cols/s43_headers。"""
    groups = p.get('s43_groups')
    if groups:
        return groups
    headers = p.get('s43_headers', [])
    cols = p.get('s43_cols', [])
    if headers or cols:
        return [{'title': p.get('s43_title', ''), 'headers': headers, 'cols': cols}]
    return []


def _s26_s27_html(s, headers, rows, per_compound, e):
    """HTML 预览：每 per_compound 行一组，每组独立表格，限度储备液浓度列用 rowspan。"""
    if not rows:
        s.append(_tbl(headers, []))
        return
    n_groups = max(1, (len(rows) + per_compound - 1) // per_compound)
    ths = "".join(f"<th>{e(str(h))}</th>" for h in headers)
    for gi in range(n_groups):
        chunk = rows[gi * per_compound: gi * per_compound + per_compound]
        if not chunk:
            continue
        limit_conc = chunk[0][1] if len(chunk[0]) > 1 else ''
        trs = []
        for ri, row in enumerate(chunk):
            cells = []
            for ci, val in enumerate(row[:len(headers)]):
                if ci == 1:  # 限度储备液浓度列
                    if ri == 0:
                        cells.append(f'<td rowspan="{len(chunk)}">{e(str(limit_conc))}</td>')
                else:
                    cells.append(f"<td>{e(str(val))}</td>")
            trs.append("<tr>" + "".join(cells) + "</tr>")
        s.append(f"<table><thead><tr>{ths}</tr></thead><tbody>{''.join(trs)}</tbody></table>")


def _linearity_table(doc, headers, s43_cols):
    """渲染检测报告 4.3 线性表：线性方程/相关系数行横向合并"""
    if not headers:
        return
    total_cols = len(headers)
    t = doc.add_table(rows=5, cols=total_cols)
    t.style = "Table Grid"
    # 表头行
    for i, h in enumerate(headers):
        _cell(t.rows[0].cells[i], h, bold=True)
    # 浓度行
    _cell(t.rows[1].cells[0], '浓度')
    for ci in range(1, total_cols):
        col = s43_cols[ci - 1] if ci - 1 < len(s43_cols) else []
        _cell(t.rows[1].cells[ci], col[0] if len(col) > 0 else '')
    # 峰面积行
    _cell(t.rows[2].cells[0], '峰面积')
    for ci in range(1, total_cols):
        col = s43_cols[ci - 1] if ci - 1 < len(s43_cols) else []
        _cell(t.rows[2].cells[ci], col[1] if len(col) > 1 else '')
    # 线性方程行（列1~末合并）
    _cell(t.rows[3].cells[0], '线性方程')
    if total_cols > 2:
        eq_cell = t.rows[3].cells[1]
        for j in range(2, total_cols):
            eq_cell = eq_cell.merge(t.rows[3].cells[j])
    val = s43_cols[0][2] if s43_cols and len(s43_cols[0]) > 2 else ''
    _cell(t.rows[3].cells[1], val)
    # 相关系数行（列1~末合并）
    _cell(t.rows[4].cells[0], '相关系数')
    if total_cols > 2:
        r_cell = t.rows[4].cells[1]
        for j in range(2, total_cols):
            r_cell = r_cell.merge(t.rows[4].cells[j])
    val = s43_cols[0][3] if s43_cols and len(s43_cols[0]) > 3 else ''
    _cell(t.rows[4].cells[1], val)
    doc.add_paragraph()


def _signature_img_html(file_id, max_width=140, max_height=60):
    img_path = _find_upload_file(file_id)
    if not img_path:
        return ""
    try:
        import base64
        img_b64 = base64.b64encode(img_path.read_bytes()).decode()
        suffix = img_path.suffix.lower().lstrip('.')
        mime = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png", "bmp": "bmp", "tiff": "tiff"}.get(suffix, "png")
        return f'<img src="data:image/{mime};base64,{img_b64}" style="max-width:{max_width}px;max-height:{max_height}px">'
    except Exception:
        return ""


def render_detection_report(output_path: str, payload: dict):
    doc = Document()
    p = payload

    # 标题（二号宋体加粗）
    _para(doc, p.get('report_title', '检测报告'), bold=True, size=_PT_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 基础信息
    _para(doc, f"文档编号：{p.get('doc_number', '')}    密级：{p.get('security_level', '')}    项目名称：{p.get('project_name', '')}")

    # 人员信息
    pers = p.get('personnel', [])
    signatures = p.get('approver_signatures') or []
    if pers:
        _table(doc, ['编制人', '审核人', '批准人', '编制日期', '版本号'],
               [[r.get('editor', ''), r.get('reviewer', ''), '[签名]' if idx < len(signatures) and signatures[idx] else r.get('approver', ''),
                 r.get('date', ''), r.get('version', '')] for idx, r in enumerate(pers)])

    # 1. 仪器与试剂
    _h1(doc, '1. 仪器与试剂')
    _h2(doc, '1.1 仪器')
    inst_h = p.get('instrument_headers', ['名称', '型号', '编号', '厂家', '有效期'])
    _table(doc, inst_h, p.get('instruments', []))

    _h2(doc, '1.2 试剂')
    _table(doc, ['名称', '批号', '级别', '厂家'], p.get('reagents', []))

    _h2(doc, '1.3 对照品与样品')
    _table(doc, ['名称', '批号', '含量', '厂家'], p.get('references', []))

    # 2. 实验步骤
    _h1(doc, '2. 实验步骤')
    _h2(doc, '2.1 各标准品母液单标及供试品配制')
    _table(doc, ['化合物名称', '纯度(%)', '称量量 mg', '定容体积 ml', '稀释剂', '母液浓度'], p.get('s21_rows', []))

    _h2(doc, '2.2 各标准品储备液单标配制')
    _table(doc, ['化合物名称', '母液浓度\nmg/ml', '移取量\nml', '定容体积\nml', '稀释剂',
                 '一级储备液\nug/ml', '移取量\nml', '定容体积\nml', '二级储备液\nng/ml'],
           p.get('s22_rows', []))

    _h2(doc, '2.3 限度储备液配制')
    _table(doc, ['化合物名称', '二级储备液\nng/ml', '移取量\nml', '定容体积\nml', '稀释剂', '限度储备液\nng/ml'],
           p.get('s23_rows', []))

    _h2(doc, '2.4 定量限单标储备液配制')
    _table(doc, ['化合物名称', '二级储备液\nng/ml', '移取量\nml', '定容体积\nml', '稀释剂', '定量限储备液\nng/ml'],
           p.get('s24_rows', []))

    _h2(doc, '2.5 定量限及检测限溶液配制')
    _table(doc, ['化合物名称', '定量限储备液\nng/ml', '移取量\nml', '定容体积\nml',
                 '定量限\nng/ml', '移取量\nml', '定容体积\nml', '检测限\nng/ml', '稀释剂'],
           p.get('s25_rows', []))

    _h2(doc, '2.6 线性溶液配制')
    _s26_s27_split(doc,
        ['名称', '限度储备液浓度\n(ng/ml)', '移取量\nml', '定容体积\nml', '线性点浓度\nng/ml', '稀释剂'],
        p.get('s26_rows', []), per_compound=5)

    _h2(doc, '2.7 回收率溶液配制')
    _s26_s27_split(doc,
        ['名称', '限度储备液浓度\n(ng/ml)', '移取量\nml', '定容体积\nml', '回收率浓度\nng/ml', '稀释剂'],
        p.get('s27_rows', []), per_compound=3)

    # 3. 实验方法
    _h1(doc, '3. 实验方法')
    ms_methods = p.get('ms_methods', [])
    if not ms_methods:
        ms_vals = p.get('ms_values', [])
        if ms_vals:
            ms_methods = [ms_vals]
    for idx, ms_vals in enumerate(ms_methods, 1):
        _h2(doc, f'3.1 质谱方法{idx}')
        ms_rows = [[_MS_LABELS[i], ms_vals[i] if i < len(ms_vals) else ''] for i in range(len(_MS_LABELS))]
        _table(doc, ['项目', '值'], ms_rows)

    hplc_methods = p.get('hplc_methods', [])
    for idx, m in enumerate(hplc_methods, 1):
        _h2(doc, f'3.2 液相方法{idx}')
        fixed = [
            ['色谱柱', m.get('column', '')],
            ['柱温', m.get('temp', '')],
            ['进样盘温度', m.get('tray_temp', '')],
            ['进样量', m.get('inj_vol', '')],
            ['切换阀', m.get('switch_valve', '')],
        ]
        _table(doc, ['项目', '值'], fixed)
        _para(doc, '流动相梯度：')
        mobile_rows = m.get('mobile_rows', [])
        if mobile_rows:
            _table(doc, ['时间(min)', '流动相 A (%)', '流动相 B(%)'], mobile_rows)
        notes = m.get('notes', '')
        if notes:
            _para(doc, f'备注：{notes}')

    # 4. 数据分析
    _h1(doc, '4. 数据分析')
    _h2(doc, '4.1 回收率分析')
    _para(doc, '回收率（%）= （测得量 - 样本含量）/ 加入量 × 100%')
    _s41_groups_table(doc, p.get('s41_groups', []))

    _h2(doc, '4.2 定量限检测限分析')
    _table(doc, ['样本名称', '浓度\n(ng/ml)', 'S/N', '样本名称', '浓度\n(ng/ml)', 'S/N'],
           p.get('s42_rows', []))

    _h2(doc, '4.3 线性分析')
    for g in _s43_groups_from_payload(p):
        g_title = g.get('title', '')
        g_headers = ['名称'] + g.get('headers', [])
        g_cols = g.get('cols', [])
        if g_title:
            _para(doc, g_title, bold=False)
        _linearity_table(doc, g_headers, g_cols)
        chart_png = _make_linearity_chart_png(g_cols)
        if chart_png:
            doc.add_paragraph()
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(io.BytesIO(chart_png), width=Inches(4.5))

    s44 = p.get('s44_content', '').strip()
    if s44:
        _h2(doc, '4.4 结果分析')
        _para(doc, s44)

    # 5. 存储路径
    _h1(doc, '5. 存储路径')
    _para(doc, p.get('storage_path', ''))

    doc.save(output_path)


def render_detection_preview(payload: dict) -> str:
    import html as _h
    e = _h.escape
    p = payload
    s = ['<style>table{border-collapse:collapse;width:100%;margin:10px 0}th,td{border:1px solid #999;padding:6px;text-align:center}th{background:#f0f0f0;font-weight:bold}h1{font-size:18px;font-weight:bold}h2{font-size:14px;font-weight:bold;margin-top:12px}</style>']

    s.append(f"<h1 style='text-align:center'>{e(p.get('report_title', '检测报告'))}</h1>")
    s.append(f"<p>文档编号：{e(p.get('doc_number', ''))} &nbsp; 密级：{e(p.get('security_level', ''))} &nbsp; 项目名称：{e(p.get('project_name', ''))}</p>")

    pers = p.get('personnel', [])
    signatures = p.get('approver_signatures') or []
    if pers:
        s.append("<table><tr><th>编制人</th><th>审核人</th><th>批准人</th><th>编制日期</th><th>版本号</th></tr>")
        for idx, r in enumerate(pers):
            approver_html = _signature_img_html(signatures[idx] if idx < len(signatures) else None) or e(r.get('approver',''))
            s.append(f"<tr><td>{e(r.get('editor',''))}</td><td>{e(r.get('reviewer',''))}</td><td>{approver_html}</td><td>{e(r.get('date',''))}</td><td>{e(r.get('version',''))}</td></tr>")
        s.append("</table>")

    s.append("<h1>1. 仪器与试剂</h1>")
    s.append("<h2>1.1 仪器</h2>")
    inst_h = p.get('instrument_headers', ['名称', '型号', '编号', '厂家', '有效期'])
    s.append(_tbl(inst_h, p.get('instruments', [])))
    s.append("<h2>1.2 试剂</h2>")
    s.append(_tbl(['名称', '批号', '级别', '厂家'], p.get('reagents', [])))
    s.append("<h2>1.3 对照品与样品</h2>")
    s.append(_tbl(['名称', '批号', '含量', '厂家'], p.get('references', [])))

    s.append("<h1>2. 实验步骤</h1>")
    s.append("<h2>2.1 各标准品母液单标及供试品配制</h2>")
    s.append(_tbl(['化合物名称', '纯度(%)', '称量量 mg', '定容体积 ml', '稀释剂', '母液浓度'], p.get('s21_rows', [])))
    s.append("<h2>2.2 各标准品储备液单标配制</h2>")
    s.append(_tbl(['化合物名称', '母液浓度', '移取量', '定容体积', '稀释剂', '一级储备液', '移取量', '定容体积', '二级储备液'], p.get('s22_rows', [])))
    s.append("<h2>2.3 限度储备液配制</h2>")
    s.append(_tbl(['化合物名称', '二级储备液', '移取量', '定容体积', '稀释剂', '限度储备液'], p.get('s23_rows', [])))
    s.append("<h2>2.4 定量限单标储备液配制</h2>")
    s.append(_tbl(['化合物名称', '二级储备液', '移取量', '定容体积', '稀释剂', '定量限储备液'], p.get('s24_rows', [])))
    s.append("<h2>2.5 定量限及检测限溶液配制</h2>")
    s.append(_tbl(['化合物名称', '定量限储备液', '移取量', '定容体积', '定量限', '移取量', '定容体积', '检测限', '稀释剂'], p.get('s25_rows', [])))
    s.append("<h2>2.6 线性溶液配制</h2>")
    _s26_s27_html(s, ['名称', '限度储备液浓度', '移取量', '定容体积', '线性点浓度', '稀释剂'], p.get('s26_rows', []), 5, e)
    s.append("<h2>2.7 回收率溶液配制</h2>")
    _s26_s27_html(s, ['名称', '限度储备液浓度', '移取量', '定容体积', '回收率浓度', '稀释剂'], p.get('s27_rows', []), 3, e)

    s.append("<h1>3. 实验方法</h1>")
    ms_methods = p.get('ms_methods', [])
    if not ms_methods:
        ms_vals = p.get('ms_values', [])
        if ms_vals:
            ms_methods = [ms_vals]
    for idx, ms_vals in enumerate(ms_methods, 1):
        s.append(f"<h2>3.1 质谱方法{idx}</h2>")
        ms_rows = [[_MS_LABELS[i], ms_vals[i] if i < len(ms_vals) else ''] for i in range(len(_MS_LABELS))]
        s.append(_tbl(['项目', '值'], ms_rows))
    for idx, m in enumerate(p.get('hplc_methods', []), 1):
        s.append(f"<h2>3.2 液相方法{idx}</h2>")
        s.append(_tbl(['项目', '值'], [
            ['色谱柱', m.get('column', '')],
            ['柱温', m.get('temp', '')],
            ['进样盘温度', m.get('tray_temp', '')],
            ['进样量', m.get('inj_vol', '')],
            ['切换阀', m.get('switch_valve', '')],
        ]))
        mobile_rows = m.get('mobile_rows', [])
        if mobile_rows:
            s.append(_tbl(['时间(min)', '流动相 A (%)', '流动相 B(%)'], mobile_rows))
        notes = m.get('notes', '')
        if notes:
            s.append(f"<p>备注：{e(notes)}</p>")

    s.append("<h1>4. 数据分析</h1>")
    s.append("<h2>4.1 回收率分析</h2>")
    s.append("<p>回收率（%）= （测得量 - 样本含量）/ 加入量 × 100%</p>")
    ths41 = "".join(f"<th>{e(h)}</th>" for h in ['样本名称', '加入量', '样本含量', '测得量', '回收率(%)', '平均回收率(%)', 'RSD(%)'])
    for g41 in p.get('s41_groups', []):
        g41_title = g41.get('title', '')
        g41_rows = g41.get('rows', [])
        if g41_title:
            s.append(f"<p><strong>{e(g41_title)}</strong></p>")
        avg_str, rsd_str = _s41_calc_avg_rsd(g41_rows)
        trs41 = []
        for ri, row in enumerate(g41_rows):
            cells = "".join(f"<td>{e(str(row[ci] if ci < len(row) else ''))}</td>" for ci in range(5))
            if ri == 0:
                cells += f'<td rowspan="{len(g41_rows)}">{e(avg_str)}</td><td rowspan="{len(g41_rows)}">{e(rsd_str)}</td>'
            trs41.append(f"<tr>{cells}</tr>")
        s.append(f"<table><thead><tr>{ths41}</tr></thead><tbody>{''.join(trs41)}</tbody></table>")
    s.append("<h2>4.2 定量限检测限分析</h2>")
    s.append(_tbl(['样本名称', '浓度', 'S/N', '样本名称', '浓度', 'S/N'], p.get('s42_rows', [])))
    s.append("<h2>4.3 线性分析</h2>")
    for g in _s43_groups_from_payload(p):
        g_title = g.get('title', '')
        g_headers = ['名称'] + g.get('headers', [])
        g_cols = g.get('cols', [])
        span = max(len(g_cols), 1)
        eq_val = g_cols[0][2] if g_cols and len(g_cols[0]) > 2 else ''
        r_val = g_cols[0][3] if g_cols and len(g_cols[0]) > 3 else ''
        if g_title:
            s.append(f"<p><strong>{e(g_title)}</strong></p>")
        rows_html = (
            "<tr><td>浓度</td>" + "".join(f"<td>{e(str(col[0] if len(col)>0 else ''))}</td>" for col in g_cols) + "</tr>"
            "<tr><td>峰面积</td>" + "".join(f"<td>{e(str(col[1] if len(col)>1 else ''))}</td>" for col in g_cols) + "</tr>"
            f"<tr><td>线性方程</td><td colspan=\"{span}\">{e(eq_val)}</td></tr>"
            f"<tr><td>相关系数</td><td colspan=\"{span}\">{e(r_val)}</td></tr>"
        )
        s.append("<table><thead><tr>" + "".join(f"<th>{e(str(h))}</th>" for h in g_headers) + "</tr></thead><tbody>" + rows_html + "</tbody></table>")

    s44 = p.get('s44_content', '').strip()
    if s44:
        s.append("<h2>4.4 结果分析</h2>")
        s.append(f"<p style='white-space:pre-wrap'>{e(s44)}</p>")

    s.append("<h1>5. 存储路径</h1>")
    s.append(f"<p>{e(p.get('storage_path', ''))}</p>")

    return ''.join(s)


def _tbl(headers, rows):
    if not headers:
        return ""
    import html as _h
    e = _h.escape
    ths = "".join(f"<th>{e(str(h))}</th>" for h in headers)
    trs = "".join("<tr>" + "".join(f"<td>{e(str(v))}</td>" for v in row) + "</tr>" for row in rows)
    return f"<table><thead><tr>{ths}</tr></thead><tbody>{trs}</tbody></table>"


def render_detection_pdf(output_path: str, payload: dict):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import glob as _glob
    import xml.sax.saxutils as _sax

    font_name = "Helvetica"
    for pat in [
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "/system/fonts/NotoSansCJK-Regular.ttc",
        "/system/fonts/DroidSansFallback.ttf",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttf",
        "/System/Library/Fonts/PingFang.ttc",
    ]:
        found = _glob.glob(pat)
        if found:
            try:
                pdfmetrics.registerFont(TTFont("CJK", found[0]))
                font_name = "CJK"
                break
            except Exception:
                continue

    cover = ParagraphStyle("cover", fontName=font_name, fontSize=18, alignment=1, spaceAfter=8, leading=24)
    h1 = ParagraphStyle("h1", fontName=font_name, fontSize=13, spaceBefore=10, spaceAfter=4, leading=18)
    h2 = ParagraphStyle("h2", fontName=font_name, fontSize=11, spaceBefore=8, spaceAfter=3, leading=16)
    body = ParagraphStyle("body", fontName=font_name, fontSize=10, spaceAfter=3, leading=14)
    cell = ParagraphStyle("cell", fontName=font_name, fontSize=9, leading=12, wordWrap='CJK')

    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2.5*cm, bottomMargin=2.5*cm)

    def _c(v):
        return Paragraph(_sax.escape(str(v or "")), cell)

    def _p(v, st):
        return Paragraph(_sax.escape(str(v or "")), st)

    def tbl(headers, rows):
        if not headers:
            return []
        data = [[_c(h) for h in headers]] + ([[_c(v) for v in row] for row in rows] if rows else [[_c("") for _ in headers]])
        col_w = doc.width / max(len(headers), 1)
        t = Table(data, colWidths=[col_w] * len(headers), repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8edf5")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        return [t, Spacer(1, 6)]

    def tbl_merged_col1(headers, rows, per_compound):
        """生成限度储备液浓度列（col 1）按组纵向合并的 PDF 表格"""
        if not headers or not rows:
            return tbl(headers, rows)
        data = [[_c(h) for h in headers]] + [[_c(v) for v in row] for row in rows]
        col_w = doc.width / max(len(headers), 1)
        t = Table(data, colWidths=[col_w] * len(headers), repeatRows=1)
        style_cmds = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8edf5")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]
        n_rows = len(rows)
        g = 0
        while g < n_rows:
            end = min(g + per_compound - 1, n_rows - 1)
            if end > g:
                style_cmds.append(("SPAN", (1, g + 1), (1, end + 1)))
            g += per_compound
        t.setStyle(TableStyle(style_cmds))
        return [t, Spacer(1, 6)]

    def tbl_split(headers, rows, per_compound):
        """每 per_compound 行一组，每组独立表（独立表头），组内 col1 纵向合并。"""
        if not rows:
            return tbl(headers, rows)
        result = []
        n_groups = max(1, (len(rows) + per_compound - 1) // per_compound)
        for gi in range(n_groups):
            chunk = rows[gi * per_compound: gi * per_compound + per_compound]
            if not chunk:
                continue
            data = [[_c(h) for h in headers]] + [[_c(v) for v in row] for row in chunk]
            col_w = doc.width / max(len(headers), 1)
            t = Table(data, colWidths=[col_w] * len(headers), repeatRows=1)
            style_cmds = [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8edf5")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
            if len(chunk) > 1:
                style_cmds.append(("SPAN", (1, 1), (1, len(chunk))))
            t.setStyle(TableStyle(style_cmds))
            result += [t, Spacer(1, 6)]
        return result

    p = payload
    story = [
        _p(p.get('report_title', '检测报告'), cover),
        _p(f"文档编号：{p.get('doc_number', '')}  密级：{p.get('security_level', '')}  项目名称：{p.get('project_name', '')}", body),
        Spacer(1, 8),
    ]

    pers = p.get('personnel', [])
    if pers:
        story += tbl(['编制人', '审核人', '批准人', '编制日期', '版本号'], [
            [r.get('editor', ''), r.get('reviewer', ''), r.get('approver', ''), r.get('date', ''), r.get('version', '')]
            for r in pers
        ])

    story += [_p('1. 仪器与试剂', h1), _p('1.1 仪器', h2)]
    story += tbl(p.get('instrument_headers', ['名称', '型号', '编号', '厂家', '有效期']), p.get('instruments', []))
    story += [_p('1.2 试剂', h2)]
    story += tbl(['名称', '批号', '级别', '厂家'], p.get('reagents', []))
    story += [_p('1.3 对照品与样品', h2)]
    story += tbl(['名称', '批号', '含量', '厂家'], p.get('references', []))

    story += [_p('2. 实验步骤', h1)]
    story += [_p('2.1 各标准品母液单标及供试品配制', h2)]
    story += tbl(['化合物名称', '纯度(%)', '称量量 mg', '定容体积 ml', '稀释剂', '母液浓度'], p.get('s21_rows', []))
    story += [_p('2.2 各标准品储备液单标配制', h2)]
    story += tbl(['化合物名称', '母液浓度', '移取量', '定容体积', '稀释剂', '一级储备液', '移取量', '定容体积', '二级储备液'], p.get('s22_rows', []))
    story += [_p('2.3 限度储备液配制', h2)]
    story += tbl(['化合物名称', '二级储备液', '移取量', '定容体积', '稀释剂', '限度储备液'], p.get('s23_rows', []))
    story += [_p('2.4 定量限单标储备液配制', h2)]
    story += tbl(['化合物名称', '二级储备液', '移取量', '定容体积', '稀释剂', '定量限储备液'], p.get('s24_rows', []))
    story += [_p('2.5 定量限及检测限溶液配制', h2)]
    story += tbl(['化合物名称', '定量限储备液', '移取量', '定容体积', '定量限', '移取量', '定容体积', '检测限', '稀释剂'], p.get('s25_rows', []))
    story += [_p('2.6 线性溶液配制', h2)]
    story += tbl_split(['名称', '限度储备液浓度', '移取量', '定容体积', '线性点浓度', '稀释剂'], p.get('s26_rows', []), per_compound=5)
    story += [_p('2.7 回收率溶液配制', h2)]
    story += tbl_split(['名称', '限度储备液浓度', '移取量', '定容体积', '回收率浓度', '稀释剂'], p.get('s27_rows', []), per_compound=3)

    story += [_p('3. 实验方法', h1)]
    ms_methods = p.get('ms_methods', [])
    if not ms_methods:
        ms_vals = p.get('ms_values', [])
        if ms_vals:
            ms_methods = [ms_vals]
    for idx, ms_vals in enumerate(ms_methods, 1):
        story += [_p(f'3.1 质谱方法{idx}', h2)]
        story += tbl(['项目', '值'], [[_MS_LABELS[i], ms_vals[i] if i < len(ms_vals) else ''] for i in range(len(_MS_LABELS))])

    for idx, m in enumerate(p.get('hplc_methods', []), 1):
        story += [_p(f'3.2 液相方法{idx}', h2)]
        story += tbl(['项目', '值'], [
            ['色谱柱', m.get('column', '')],
            ['柱温', m.get('temp', '')],
            ['进样盘温度', m.get('tray_temp', '')],
            ['进样量', m.get('inj_vol', '')],
            ['切换阀', m.get('switch_valve', '')],
        ])
        mobile_rows = m.get('mobile_rows', [])
        if mobile_rows:
            story += tbl(['时间(min)', '流动相 A (%)', '流动相 B(%)'], mobile_rows)
        if m.get('notes'):
            story += [_p(f"备注：{m.get('notes')}", body)]

    story += [_p('4. 数据分析', h1), _p('4.1 回收率分析', h2)]
    story += [_p('回收率（%）= （测得量 - 样本含量）/ 加入量 × 100%', body)]
    s41_headers_pdf = ['样本名称', '加入量', '样本含量', '测得量', '回收率(%)', '平均回收率(%)', 'RSD(%)']
    for g41 in p.get('s41_groups', []):
        g41_title = g41.get('title', '')
        g41_rows = g41.get('rows', [])
        if g41_title:
            story += [_p(g41_title, body)]
        if not g41_rows:
            continue
        avg_str, rsd_str = _s41_calc_avg_rsd(g41_rows)
        col_w = doc.width / 7
        data = [[_c(h) for h in s41_headers_pdf]]
        for ri, row in enumerate(g41_rows):
            data.append([_c(row[ci] if ci < len(row) else '') for ci in range(5)] + [_c(avg_str if ri == 0 else ''), _c(rsd_str if ri == 0 else '')])
        from reportlab.platypus import Table as _RLTbl, TableStyle as _RLStyle
        from reportlab.lib import colors as _rlc
        t41 = _RLTbl(data, colWidths=[col_w] * 7, repeatRows=1)
        style_cmds = [
            ("BACKGROUND", (0, 0), (-1, 0), _rlc.HexColor("#e8edf5")),
            ("GRID", (0, 0), (-1, -1), 0.5, _rlc.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]
        if len(g41_rows) > 1:
            style_cmds += [
                ("SPAN", (5, 1), (5, len(g41_rows))),
                ("SPAN", (6, 1), (6, len(g41_rows))),
            ]
        t41.setStyle(_RLStyle(style_cmds))
        story += [t41, Spacer(1, 6)]
    story += [_p('4.2 定量限检测限分析', h2)]
    story += tbl(['样本名称', '浓度', 'S/N', '样本名称', '浓度', 'S/N'], p.get('s42_rows', []))
    story += [_p('4.3 线性分析', h2)]
    for g in _s43_groups_from_payload(p):
        g_title = g.get('title', '')
        g_headers = ['名称'] + g.get('headers', [])
        g_cols = g.get('cols', [])
        if g_title:
            story += [_p(g_title, body)]
        g_data = [
            ['浓度'] + [col[0] if len(col) > 0 else '' for col in g_cols],
            ['峰面积'] + [col[1] if len(col) > 1 else '' for col in g_cols],
        ]
        story += tbl(g_headers, g_data)
        eq_val = g_cols[0][2] if g_cols and len(g_cols[0]) > 2 else ''
        r_val = g_cols[0][3] if g_cols and len(g_cols[0]) > 3 else ''
        from reportlab.platypus import Table as RLTable, TableStyle as RLTableStyle
        from reportlab.lib import colors as rcolors
        col_w1 = doc.width / max(len(g_headers), 1)
        col_w2 = doc.width - col_w1
        for label, val in [('线性方程', eq_val), ('相关系数', r_val)]:
            t_s = RLTable([[_c(label), _c(val)]], colWidths=[col_w1, col_w2])
            t_s.setStyle(RLTableStyle([('GRID', (0,0),(-1,-1),0.5,rcolors.grey), ('VALIGN',(0,0),(-1,-1),'MIDDLE')]))
            story += [t_s]
        chart_png = _make_linearity_chart_png(g_cols)
        if chart_png:
            from reportlab.platypus import Image as RLImage
            story += [Spacer(1, 6), RLImage(io.BytesIO(chart_png), width=doc.width * 0.7, height=doc.width * 0.7 * 3.5 / 5)]

    s44 = p.get('s44_content', '').strip()
    if s44:
        story += [_p('4.4 结果分析', h2), _p(s44, body)]

    story += [_p('5. 存储路径', h1), _p(p.get('storage_path', ''), body)]

    def _footer(canvas, d):
        canvas.saveState()
        canvas.setFont(font_name, 9)
        canvas.drawCentredString(A4[0] / 2, 1.2 * cm, f"第 {d.page} 页")
        canvas.restoreState()

    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)

