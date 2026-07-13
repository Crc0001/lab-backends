from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

REPORTS_DIR = Path("static/reports")

_SONG = "宋体"
_TNR = "Times New Roman"
_PT_BODY = Pt(10.5)   # 小四：正文
_PT_TABLE = Pt(9)     # 五号：表格内容
_PT_H1 = Pt(12)       # 三号：一级标题


def _set_run_font(run, bold=False, size=None):
    run.bold = bold
    if size:
        run.font.size = size
    run.font.name = _TNR
    run._element.rPr.rFonts.set(qn("w:eastAsia"), _SONG)


def _add_h1(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, bold=True, size=_PT_H1)


def _add_h2(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, bold=True, size=_PT_BODY)


def _set_cell_text(cell, text: str, bold=False):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    cell.paragraphs[0].clear()
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = cell.paragraphs[0].add_run(str(text))
    _set_run_font(run, bold=bold, size=_PT_TABLE)


def _add_table(doc, headers: list, rows: list):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row[:len(headers)]):
            _set_cell_text(table.rows[r_idx + 1].cells[c_idx], val)
    doc.add_paragraph()


def _merge_cells_vertical(table, col: int, start_row: int, end_row: int):
    a = table.cell(start_row, col)
    b = table.cell(end_row, col)
    a.merge(b)





def _build_conditions_table(doc, p):
    """
    构建检测条件多级表（含合并单元格）。
    列布局：col0=分组标签, col1=检测条件, col2=确定依据
    行布局（从上到下）：
      0: 表头
      1-7: 液相条件（色谱柱/柱温/流动相+梯度/流速/进样体积/进样器温度）
      8-12: 质谱条件（质谱仪/电离源/扫描模式/监测模式+子表/离子源参数+子表）
      13: 限度
      14: 切换阀
      15: 检测时间
    """
    # 计算总行数：
    # 表头1 + 液相(色谱柱1+柱温1+流动相desc1+梯度header1+梯度rows+流速1+进样体积1+进样器温度1)
    # + 质谱(质谱仪1+电离源1+扫描模式1+监测模式header1+监测rows+离子源header1+离子源rows)
    # + 限度1+切换阀1+检测时间1
    mp_rows = p.mobile_phase_rows or [["", "", ""]]
    mon_rows = p.monitoring_rows or [["", "", ""]]
    ion_rows = p.ion_source_rows or [["", "", "", "", ""]]

    # 行索引规划
    ROW_HEADER = 0
    # 液相块
    ROW_HPLC_START = 1
    ROW_HPLC_COL = 1        # 色谱柱
    ROW_HPLC_TEMP = 2       # 柱温
    ROW_HPLC_MP_DESC = 3    # 流动相描述
    ROW_HPLC_MP_HEAD = 4    # 梯度表头
    ROW_HPLC_MP_DATA = 5    # 梯度数据起始
    hplc_mp_data_end = ROW_HPLC_MP_DATA + len(mp_rows) - 1
    ROW_HPLC_FLOW = hplc_mp_data_end + 1
    ROW_HPLC_INJ_VOL = ROW_HPLC_FLOW + 1
    ROW_HPLC_INJ_TEMP = ROW_HPLC_INJ_VOL + 1
    ROW_HPLC_END = ROW_HPLC_INJ_TEMP
    # 质谱块
    ROW_MS_START = ROW_HPLC_END + 1
    ROW_MS_INST = ROW_MS_START
    ROW_MS_ION = ROW_MS_INST + 1
    ROW_MS_SCAN = ROW_MS_ION + 1
    ROW_MS_MON_HEAD = ROW_MS_SCAN + 1
    ROW_MS_MON_DATA = ROW_MS_MON_HEAD + 1
    ms_mon_data_end = ROW_MS_MON_DATA + len(mon_rows) - 1
    ROW_MS_ION_HEAD = ms_mon_data_end + 1
    ROW_MS_ION_DATA = ROW_MS_ION_HEAD + 1
    ms_ion_data_end = ROW_MS_ION_DATA + len(ion_rows) - 1
    ROW_MS_END = ms_ion_data_end
    # 底部
    ROW_LIMIT = ROW_MS_END + 1
    ROW_SWITCH = ROW_LIMIT + 1
    ROW_DET_TIME = ROW_SWITCH + 1
    TOTAL_ROWS = ROW_DET_TIME + 1

    table = doc.add_table(rows=TOTAL_ROWS, cols=3)
    table.style = "Table Grid"

    # ── 表头 ──
    _set_cell_text(table.cell(ROW_HEADER, 0), "/", bold=True)
    _set_cell_text(table.cell(ROW_HEADER, 1), "检测条件", bold=True)
    _set_cell_text(table.cell(ROW_HEADER, 2), "确定依据", bold=True)

    # ── 液相条件块 ──
    # col0 合并整个液相块
    _merge_cells_vertical(table, 0, ROW_HPLC_START, ROW_HPLC_END)
    _set_cell_text(table.cell(ROW_HPLC_START, 0), "液相条件", bold=True)

    _set_cell_text(table.cell(ROW_HPLC_COL, 1), f"色谱柱：{p.hplc_column}")
    _set_cell_text(table.cell(ROW_HPLC_COL, 2), p.hplc_column_det)

    _set_cell_text(table.cell(ROW_HPLC_TEMP, 1), f"柱温：{p.hplc_temp}")
    _set_cell_text(table.cell(ROW_HPLC_TEMP, 2), p.hplc_temp_det)

    # 流动相描述
    _set_cell_text(table.cell(ROW_HPLC_MP_DESC, 1), p.hplc_mobile_phase_desc)
    # 流动相确定依据跨 desc+head+data 行
    _merge_cells_vertical(table, 2, ROW_HPLC_MP_DESC, hplc_mp_data_end)
    _set_cell_text(table.cell(ROW_HPLC_MP_DESC, 2), p.hplc_column_det)

    # 梯度表头（嵌入 col1，横向合并3子列用文本模拟）
    _set_cell_text(table.cell(ROW_HPLC_MP_HEAD, 1), "时间 | A(%) | B(%)", bold=True)
    for i, row in enumerate(mp_rows):
        _set_cell_text(table.cell(ROW_HPLC_MP_DATA + i, 1),
                       " | ".join(v for v in row))

    _set_cell_text(table.cell(ROW_HPLC_FLOW, 1), f"流速：{p.hplc_flow_rate}")
    _set_cell_text(table.cell(ROW_HPLC_FLOW, 2), p.hplc_flow_rate_det)

    _set_cell_text(table.cell(ROW_HPLC_INJ_VOL, 1), f"进样体积：{p.hplc_inj_vol}")
    _set_cell_text(table.cell(ROW_HPLC_INJ_VOL, 2), p.hplc_inj_vol_det)

    _set_cell_text(table.cell(ROW_HPLC_INJ_TEMP, 1), f"进样器温度：{p.hplc_inj_temp}")
    _set_cell_text(table.cell(ROW_HPLC_INJ_TEMP, 2), p.hplc_inj_temp_det)

    # ── 质谱条件块 ──
    _merge_cells_vertical(table, 0, ROW_MS_START, ROW_MS_END)
    _set_cell_text(table.cell(ROW_MS_START, 0), "质谱条件", bold=True)

    _set_cell_text(table.cell(ROW_MS_INST, 1), f"质谱仪：{p.ms_instrument}")
    _set_cell_text(table.cell(ROW_MS_INST, 2), p.ms_instrument_det)

    _set_cell_text(table.cell(ROW_MS_ION, 1), f"电离源：{p.ms_ion_source}")
    _set_cell_text(table.cell(ROW_MS_ION, 2), p.ms_ion_source_det)

    _set_cell_text(table.cell(ROW_MS_SCAN, 1), f"扫描模式：{p.ms_scan_mode}")
    _set_cell_text(table.cell(ROW_MS_SCAN, 2), p.ms_scan_mode_det)

    # 监测模式
    _set_cell_text(table.cell(ROW_MS_MON_HEAD, 1), "化合物 | 相对分子质量 | 母离子(+H)", bold=True)
    _merge_cells_vertical(table, 2, ROW_MS_MON_HEAD, ms_mon_data_end)
    _set_cell_text(table.cell(ROW_MS_MON_HEAD, 2), p.ms_instrument_det)
    for i, row in enumerate(mon_rows):
        _set_cell_text(table.cell(ROW_MS_MON_DATA + i, 1), " | ".join(v for v in row))

    # 离子源参数
    ion_headers = "喷雾电压(KV) | 鞘气 | 辅助气 | 离子源温度(°C) | 去溶气温度(°C)"
    _set_cell_text(table.cell(ROW_MS_ION_HEAD, 1), ion_headers, bold=True)
    _merge_cells_vertical(table, 2, ROW_MS_ION_HEAD, ms_ion_data_end)
    _set_cell_text(table.cell(ROW_MS_ION_HEAD, 2), p.ion_source_det)
    for i, row in enumerate(ion_rows):
        _set_cell_text(table.cell(ROW_MS_ION_DATA + i, 1), " | ".join(v for v in row))

    # ── 底部单行 ──
    _set_cell_text(table.cell(ROW_LIMIT, 0), "限度")
    _set_cell_text(table.cell(ROW_LIMIT, 1), p.limit)
    _set_cell_text(table.cell(ROW_LIMIT, 2), p.limit_det)

    _set_cell_text(table.cell(ROW_SWITCH, 0), "切换阀")
    _set_cell_text(table.cell(ROW_SWITCH, 1), p.switch_valve)
    _set_cell_text(table.cell(ROW_SWITCH, 2), p.switch_valve_det)

    _set_cell_text(table.cell(ROW_DET_TIME, 0), "检测时间")
    _set_cell_text(table.cell(ROW_DET_TIME, 1), p.detection_time)
    _set_cell_text(table.cell(ROW_DET_TIME, 2), p.detection_time_det)

    doc.add_paragraph()


def render_report(output_path: str, payload) -> None:
    doc = Document()

    # ── 封面 ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(payload.report_title or "化学检测实验报告")
    _set_run_font(run, bold=True, size=Pt(16))

    for label, val in [
        ("文档编号", payload.doc_number),
        ("密级", payload.security_level),
        ("项目名称", payload.project_name),
        ("生成日期", datetime.now().strftime("%Y-%m-%d")),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}：{val}")
        _set_run_font(run, size=_PT_BODY)
    doc.add_paragraph()

    # ── 1. 检测方法 ──
    _add_h1(doc, "1. 检测方法")
    if payload.detection_method:
        p = doc.add_paragraph()
        run = p.add_run(payload.detection_method)
        _set_run_font(run, size=_PT_BODY)
        doc.add_paragraph()

    _add_h2(doc, "杂质限度表")
    _add_table(doc,
               ["杂质名称", "每日最大允许摄入量", "每日最大摄入量", "限度", "供试品浓度", "限度浓度"],
               payload.impurities)

    _add_h2(doc, "检测条件表")
    _build_conditions_table(doc, payload)

    if payload.detection_note:
        p = doc.add_paragraph()
        run = p.add_run(payload.detection_note)
        _set_run_font(run, size=_PT_BODY)
        doc.add_paragraph()

    # ── 2. 方法学验证标准 ──
    _add_h1(doc, "2. 方法学验证标准")
    p = doc.add_paragraph()
    run = p.add_run(
        "本分析方法要验证的项目有：专属性、定量限和检测限、准确度重复性、线性与范围、样品检测。"
    )
    _set_run_font(run, size=_PT_BODY)
    doc.add_paragraph()
    _add_table(doc, ["项目", "可接受标准", "结果"], payload.validation_rows)

    # ── 3. 仪器与试剂 ──
    _add_h1(doc, "3. 仪器与试剂")
    _add_h2(doc, "3.1 仪器")
    _add_table(doc, ["名称", "型号", "编号", "厂家", "有效期"], payload.instruments)
    _add_h2(doc, "3.2 试剂")
    _add_table(doc, ["名称", "批号", "级别", "厂家"], payload.reagents)
    _add_h2(doc, "3.3 对照品与样品")
    _add_table(doc, ["名称", "批号", "含量（%）", "厂家"], payload.references)

    # ── 4. 验证内容 ──
    if payload.chapter4:
        _render_chapter4(doc, payload.chapter4)

    # ── 5. 修订历史 ──
    _add_h1(doc, "5. 修订历史")
    _add_table(doc, ["文件编号", "更改原因", "生效日期"], payload.revision_rows or [["", "", ""]])

    # ── 附表1：人员培训表 ──
    doc.add_page_break()
    _add_h1(doc, "附表1：人员培训表")
    _build_training_table(doc, payload)

    # ── 附表2：偏差处理表 ──
    doc.add_page_break()
    _add_h1(doc, "附表2：偏差处理表")
    _add_table(doc, ["偏差编号", "偏差名称", "记录人/日期", "备注"],
               payload.deviation_rows or [["", "", "", ""]])

    doc.save(output_path)


def _build_recovery_rows(tbl_def):
    if not tbl_def:
        return []
    if getattr(tbl_def, 'rows', None):
        return tbl_def.rows or []
    sub_rows = getattr(tbl_def, 'sub_rows', None) or []
    rows = []
    for i, sub in enumerate(sub_rows):
        if len(sub) >= 5:
            rows.append([
                sub[0] if len(sub) > 0 else '',
                sub[1] if len(sub) > 1 else '',
                sub[2] if len(sub) > 2 else '',
                sub[3] if len(sub) > 3 else '',
                sub[4] if len(sub) > 4 else '',
                getattr(tbl_def, 'avg_recovery', '') if i == 0 else '',
                getattr(tbl_def, 'rsd', '') if i == 0 else '',
            ])
        else:
            rows.append([
                getattr(tbl_def, 'group_label', '') if i == 0 else '',
                getattr(tbl_def, 'added_amount', '') if i == 0 else '',
                getattr(tbl_def, 'sample_content', '') if i == 0 else '',
                sub[1] if len(sub) > 1 else '',
                sub[2] if len(sub) > 2 else '',
                getattr(tbl_def, 'avg_recovery', '') if i == 0 else '',
                getattr(tbl_def, 'rsd', '') if i == 0 else '',
            ])
    return rows


def _build_recovery_headers(tbl_def):
    headers = getattr(tbl_def, 'headers', None) or []
    return headers or ["大类名称 / 子名称", "加入量(ng/ml)", "样本含量(ng/ml)", "测得量(ng/ml)", "回收率(%)", "平均回收率(%)", "RSD(%)"]


def _build_recovery_table(doc, tbl_def):
    headers = _build_recovery_headers(tbl_def)
    rows = _build_recovery_rows(tbl_def)
    if not rows:
        rows = [["", "", "", "", "", "", ""] for _ in range(6)]

    total_rows = 1 + len(rows)
    total_cols = len(headers)
    table = doc.add_table(rows=total_rows, cols=total_cols)
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True)

    for row_idx, row in enumerate(rows, start=1):
        for col_idx in range(min(len(row), total_cols)):
            _set_cell_text(table.rows[row_idx].cells[col_idx], row[col_idx])

    if len(rows) > 1:
        for col in [5, 6]:
            _merge_cells_vertical(table, col, 1, len(rows))

    doc.add_paragraph()


# ── Chapter 4 helpers ─────────────────────────────────────────────────────────


def _render_chapter4(doc, c4):
    compounds = c4.compounds or []
    names = [c.name for c in compounds]

    _add_h1(doc, "4. 验证内容")

    # ── 4.1 专属性 ────────────────────────────────────────────────────────────
    _add_h2(doc, "4.1 专属性")
    for i, text in enumerate(c4.s41_items):
        _body(doc, f"（{i + 1}）{text}" if not text.startswith('（') else text)
    _body(doc, c4.s41_table_title or "表4.1-1  亚硝胺杂质专属性实验结果")
    _add_editable_table(doc, c4.s41_table,
                        default_headers=["杂质名称", "空白溶剂\n保留时间min", "限度定位溶液（线性3）\n保留时间min", "供试品溶液\n保留时间min"],
                        empty_rows=len(compounds) or 1)
    _body(doc, f"结论：{c4.s41_conclusion}")
    _body(doc, "")

    # ── 4.2 定量限和检测限 ────────────────────────────────────────────────────
    _add_h2(doc, "4.2 定量限和检测限")
    for i, text in enumerate(c4.s42_items):
        _body(doc, f"（{i + 1}）{text}" if not text.startswith('（') else text)
    loq_default = ["序号", "保留时间(min)", "峰面积", "S/N", "浓度(ng/ml)", "相当于样品的量(ppm)"]
    for i, name in enumerate(names):
        title = c4.s42_loq_titles[i] if i < len(c4.s42_loq_titles) else f"表4.2-{i+1} {name}定量限实验结果"
        _body(doc, title)
        tbl = c4.s42_loq_tables[i] if i < len(c4.s42_loq_tables) else None
        _add_editable_table(doc, tbl, default_headers=loq_default, empty_rows=5)
    _body(doc, c4.s42_lod_title or f"表4.2-{len(compounds)+1} 亚硝胺杂质检测限实验结果")
    lod_default = ["名称", "保留时间(min)", "峰面积", "S/N", "浓度(ng/ml)", "相当于样品的量(ppm)"]
    _add_editable_table(doc, c4.s42_lod_table, default_headers=lod_default, empty_rows=len(compounds) or 1)
    _body(doc, f"结论：{c4.s42_conclusion}")
    _body(doc, "")

    # ── 4.3 线性与范围 ────────────────────────────────────────────────────────
    _add_h2(doc, "4.3 线性与范围")
    for i, text in enumerate(c4.s43_items):
        _body(doc, f"（{i + 1}）{text}" if not text.startswith('（') else text)
    for i, name in enumerate(names):
        title = c4.s43_titles[i] if i < len(c4.s43_titles) else f"表4.3-{i+1} {name}线性实验结果"
        _body(doc, title)
        tbl = c4.s43_tables[i] if i < len(c4.s43_tables) else None
        _build_linearity_table(doc, tbl)
        _body(doc, f"图4.3-{i+1} {name}线性关系图")
        _body(doc, "")
    _body(doc, f"结论：{c4.s43_conclusion}")
    _body(doc, "")

    # ── 4.4 准确度和重复性 ────────────────────────────────────────────────────
    _add_h2(doc, "4.4 准确度和重复性")
    for i, text in enumerate(c4.s44_items):
        _body(doc, f"（{i + 1}）{text}" if not text.startswith('（') else text)
    acc_default = ["NDMA", "加入量(ng/ml)", "样本含量(ng/ml)", "测得量(ng/ml)", "回收率(%)", "平均回收率(%)", "RSD(%)"]
    for i, name in enumerate(names):
        title = c4.s44_titles[i] if i < len(c4.s44_titles) else f"表4.4-{i+1} {name}回收率实验结果"
        _body(doc, title)
        tbl = c4.s44_tables[i] if i < len(c4.s44_tables) else None
        if tbl and getattr(tbl, 'sub_rows', None) is not None:
            _build_recovery_table(doc, tbl)
        else:
            _add_editable_table(doc, tbl, default_headers=acc_default, empty_rows=6)
    _body(doc, f"结论：{c4.s44_conclusion}")
    _body(doc, "")

    # ── 4.5 样品检测 ─────────────────────────────────────────────────────────
    _add_h2(doc, "4.5 样品检测")
    for i, text in enumerate(c4.s45_items):
        _body(doc, f"（{i + 1}）{text}" if not text.startswith('（') else text)
    _body(doc, c4.s45_table_title or "表4.5 亚硝胺杂质检测结果")
    _build_s45_table(doc, c4.s45_table, names)
    _body(doc, f"结论：{c4.s45_conclusion}")
    _body(doc, "")


def _body(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size=_PT_BODY)


def _add_editable_table(doc, tbl_def, default_headers: list, empty_rows: int):
    """渲染表格：有用户数据则输出数据行，否则输出空行占位。"""
    headers = (tbl_def.headers if tbl_def and tbl_def.headers else default_headers)
    rows = (tbl_def.rows if tbl_def and tbl_def.rows else [])
    if not rows:
        rows = [[""] * len(headers) for _ in range(empty_rows)]
    cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            val = row[c_idx] if c_idx < len(row) else ""
            _set_cell_text(table.rows[r_idx + 1].cells[c_idx], val)
    doc.add_paragraph()


def _build_linearity_table(doc, tbl):
    """
    渲染 4.3 线性表，结构：
      列0: 名称（固定行标签）
      列1~6: 定量限均值 + S1~S5（可编辑表头）
    行：表头 / 浓度 / 峰面积 / 线性方程（跨列）/ 相关系数（跨列）
    """
    default_col_headers = ["定量限均值", "20XX-S1", "20XX-S2", "20XX-S3", "20XX-S4", "20XX-S5"]
    col_headers = (tbl.col_headers if tbl and tbl.col_headers else default_col_headers)
    concentrations = (tbl.concentrations if tbl and tbl.concentrations else [""] * len(col_headers))
    peak_areas = (tbl.peak_areas if tbl and tbl.peak_areas else [""] * len(col_headers))
    equation = tbl.equation if tbl else ""
    r2 = tbl.r2 if tbl else ""

    n_data_cols = len(col_headers)
    total_cols = 1 + n_data_cols  # 名称列 + 数据列

    table = doc.add_table(rows=5, cols=total_cols)
    table.style = "Table Grid"

    # 表头行
    _set_cell_text(table.rows[0].cells[0], "名称", bold=True)
    for j, h in enumerate(col_headers):
        _set_cell_text(table.rows[0].cells[j + 1], h, bold=True)

    # 浓度行
    _set_cell_text(table.rows[1].cells[0], f"浓度(ng/ml)")
    for j, v in enumerate(concentrations[:n_data_cols]):
        _set_cell_text(table.rows[1].cells[j + 1], v)

    # 峰面积行
    _set_cell_text(table.rows[2].cells[0], "峰面积")
    for j, v in enumerate(peak_areas[:n_data_cols]):
        _set_cell_text(table.rows[2].cells[j + 1], v)

    # 线性方程行（列1~末 合并）
    _set_cell_text(table.rows[3].cells[0], "线性方程")
    eq_cell = table.rows[3].cells[1]
    for j in range(2, total_cols):
        eq_cell = eq_cell.merge(table.rows[3].cells[j])
    _set_cell_text(eq_cell, equation)

    # 相关系数行（列1~末 合并）
    _set_cell_text(table.rows[4].cells[0], "相关系数")
    r2_cell = table.rows[4].cells[1]
    for j in range(2, total_cols):
        r2_cell = r2_cell.merge(table.rows[4].cells[j])
    _set_cell_text(r2_cell, r2)

    doc.add_paragraph()


def _build_training_table(doc, payload):
    """
    附表1 人员培训表，结构（参照图片）：
      行0: 培训内容 | [值，跨3列]
      行1: 主讲人   | [值] | 培训日期 | [值]
      行2: 参加人员签到：[多行文本，跨全列]
    """
    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"

    # 行0：培训内容（列0）| 值（列1~3 合并）
    _set_cell_text(table.rows[0].cells[0], "培训内容")
    content_cell = table.rows[0].cells[1]
    content_cell = content_cell.merge(table.rows[0].cells[2])
    content_cell = content_cell.merge(table.rows[0].cells[3])
    _set_cell_text(content_cell, payload.training_content)

    # 行1：主讲人 | 值 | 培训日期 | 值
    _set_cell_text(table.rows[1].cells[0], "主讲人")
    _set_cell_text(table.rows[1].cells[1], payload.training_instructor)
    _set_cell_text(table.rows[1].cells[2], "培训日期")
    _set_cell_text(table.rows[1].cells[3], payload.training_date)

    # 行2：参加人员签到（跨全列）
    attendees_cell = table.rows[2].cells[0]
    for j in range(1, 4):
        attendees_cell = attendees_cell.merge(table.rows[2].cells[j])
    attendees_text = f"参加人员签到：\n{payload.training_attendees}" if payload.training_attendees else "参加人员签到："
    _set_cell_text(attendees_cell, attendees_text)

    doc.add_paragraph()


def _build_s45_table(doc, tbl_def, names: list):
    headers = (tbl_def.headers if tbl_def and getattr(tbl_def, 'headers', None) else [])
    raw_rows = (tbl_def.rows if tbl_def and tbl_def.rows else [])
    if not names:
        names = [f"对照品{i+1}" for i in range(max(1, (max(len(headers), 2) - 2) // 2))]
    expected_cols = 2 + len(names) * 2
    if len(raw_rows) % 2 != 0:
        raw_rows = raw_rows + [[""] * expected_cols]
    if not raw_rows:
        raw_rows = [[""] * expected_cols for _ in range(2)]

    def _header(idx, fallback):
        return headers[idx] if idx < len(headers) and headers[idx] else fallback

    groups = [list(range(i, min(i + 2, len(names)))) for i in range(0, len(names), 2)] or [[0]]
    for group in groups:
        total_cols = 2 + len(group) * 2
        table = doc.add_table(rows=2 + len(raw_rows), cols=total_cols)
        table.style = "Table Grid"

        _set_cell_text(table.cell(0, 0), _header(0, "批号"), bold=True)
        table.cell(0, 0).merge(table.cell(1, 0))
        _set_cell_text(table.cell(0, 1), _header(1, "称样量\nmg"), bold=True)
        table.cell(0, 1).merge(table.cell(1, 1))

        for local_idx, compound_idx in enumerate(group):
            base = 2 + local_idx * 2
            _set_cell_text(table.cell(0, base), names[compound_idx] or f"对照品{compound_idx + 1}", bold=True)
            table.cell(0, base).merge(table.cell(0, base + 1))
            _set_cell_text(table.cell(1, base), _header(2 + compound_idx * 2, "检测值(ppm)"), bold=True)
            _set_cell_text(table.cell(1, base + 1), _header(3 + compound_idx * 2, "均值(ppm)"), bold=True)

        for pair in range(len(raw_rows) // 2):
            r0 = 2 + pair * 2
            r1 = r0 + 1
            row0 = raw_rows[pair * 2] if pair * 2 < len(raw_rows) else []
            row1 = raw_rows[pair * 2 + 1] if pair * 2 + 1 < len(raw_rows) else []

            def _val(row, idx):
                return row[idx] if idx < len(row) else ""

            _set_cell_text(table.cell(r0, 0), _val(row0, 0))
            table.cell(r0, 0).merge(table.cell(r1, 0))
            _set_cell_text(table.cell(r0, 1), _val(row0, 1))
            _set_cell_text(table.cell(r1, 1), _val(row1, 1))

            for local_idx, compound_idx in enumerate(group):
                base = 2 + local_idx * 2
                source_base = 2 + compound_idx * 2
                _set_cell_text(table.cell(r0, base), _val(row0, source_base))
                _set_cell_text(table.cell(r1, base), _val(row1, source_base))
                _set_cell_text(table.cell(r0, base + 1), _val(row0, source_base + 1))
                table.cell(r0, base + 1).merge(table.cell(r1, base + 1))

        doc.add_paragraph()


# ── HTML 预览 ─────────────────────────────────────────────────────────────────

def render_preview_html(payload) -> str:
    import html as _html
    e = _html.escape

    def tbl(headers, rows):
        if not headers:
            return ""
        ths = "".join(f"<th>{e(str(h))}</th>" for h in headers)
        trs = "".join("<tr>" + "".join(f"<td>{e(str(v))}</td>" for v in row) + "</tr>" for row in rows)
        return f"<table><thead><tr>{ths}</tr></thead><tbody>{trs}</tbody></table>"

    def recovery_tbl(tbl_def):
        headers = _build_recovery_headers(tbl_def)
        rows = _build_recovery_rows(tbl_def)
        if not rows:
            rows = [["", "", "", "", "", "", ""] for _ in range(6)]
        rows_html = ["<tr>" + "".join(f"<th>{e(str(h))}</th>" for h in headers) + "</tr>"]
        span = len(rows)
        for i, row in enumerate(rows):
            cells = []
            for col_idx in range(5):
                cells.append(f"<td>{e(row[col_idx] if col_idx < len(row) else '')}</td>")
            if i == 0:
                cells.append(f"<td rowspan=\"{span}\">{e(row[5] if len(row) > 5 else '')}</td>")
                cells.append(f"<td rowspan=\"{span}\">{e(row[6] if len(row) > 6 else '')}</td>")
            rows_html.append("<tr>" + "".join(cells) + "</tr>")
        return "<table><tbody>" + "".join(rows_html) + "</tbody></table>"

    p = payload
    c4 = p.chapter4
    names = [c.name for c in (c4.compounds or [])]
    s = []

    s.append(f"<h1>{e(p.report_title or '化学检测实验报告')}</h1>")
    s.append(f"<p>文档编号：{e(p.doc_number)} &nbsp; 密级：{e(p.security_level)} &nbsp; 项目名称：{e(p.project_name)}</p>")
    s.append(tbl(["编制人","审核人","批准人","编制日期","版本号"],
                 [[r.editor,r.reviewer,r.approver,r.date,r.version] for r in (p.personnel or [])]))

    s.append("<h2>1. 检测方法</h2>")
    if p.detection_method:
        s.append(f"<p>{e(p.detection_method)}</p>")
    s.append("<h3>杂质限度表</h3>")
    s.append(tbl(["杂质名称","每日最大允许摄入量","每日最大摄入量","限度","供试品浓度","限度浓度"], p.impurities or []))
    mp_rows = p.mobile_phase_rows or []
    mon_rows = p.monitoring_rows or []
    ion_rows = p.ion_source_rows or []
    cond_rows = (
        [["液相条件", f"色谱柱：{p.hplc_column}", p.hplc_column_det],
         ["", f"柱温：{p.hplc_temp}", p.hplc_temp_det],
         ["", p.hplc_mobile_phase_desc, ""]]
        + [["", " | ".join(r), ""] for r in mp_rows]
        + [["", f"流速：{p.hplc_flow_rate}", p.hplc_flow_rate_det],
           ["", f"进样体积：{p.hplc_inj_vol}", p.hplc_inj_vol_det],
           ["", f"进样器温度：{p.hplc_inj_temp}", p.hplc_inj_temp_det],
           ["质谱条件", f"质谱仪：{p.ms_instrument}", p.ms_instrument_det],
           ["", f"电离源：{p.ms_ion_source}", p.ms_ion_source_det],
           ["", f"扫描模式：{p.ms_scan_mode}", p.ms_scan_mode_det]]
        + [["", " | ".join(r), ""] for r in mon_rows]
        + [["", " | ".join(r), ""] for r in ion_rows]
        + [["限度", p.limit, p.limit_det],
           ["切换阀", p.switch_valve, p.switch_valve_det],
           ["检测时间", p.detection_time, p.detection_time_det]]
    )
    s.append("<h3>检测条件表</h3>")
    s.append(tbl(["/","检测条件","确定依据"], cond_rows))

    s.append("<h2>2. 方法学验证标准</h2>")
    s.append(tbl(["项目","可接受标准","结果"], p.validation_rows or []))

    s.append("<h2>3. 仪器与试剂</h2>")
    s.append(tbl(["名称","型号","编号","厂家","有效期"], p.instruments or []))
    s.append(tbl(["名称","批号","级别","厂家"], p.reagents or []))
    s.append(tbl(["名称","批号","含量（%）","厂家"], p.references or []))

    s.append("<h2>4. 验证内容</h2>")
    s.append(f"<h3>4.1 专属性</h3><p>{e(c4.s41_table_title)}</p>")
    h41 = c4.s41_table.headers or ["杂质名称","空白溶剂保留时间min","限度定位溶液保留时间min","供试品溶液保留时间min"]
    s.append(tbl(h41, c4.s41_table.rows or []))
    s.append(f"<p>结论：{e(c4.s41_conclusion)}</p>")

    s.append("<h3>4.2 定量限和检测限</h3>")
    for i, name in enumerate(names):
        title = c4.s42_loq_titles[i] if i < len(c4.s42_loq_titles) else f"表4.2-{i+1} {name}定量限"
        td = c4.s42_loq_tables[i] if i < len(c4.s42_loq_tables) else None
        s.append(f"<p>{e(title)}</p>")
        s.append(tbl(td.headers if td and td.headers else ["序号","保留时间(min)","峰面积","S/N","浓度(ng/ml)","相当于样品的量(ppm)"],
                     td.rows if td else []))
    s.append(f"<p>{e(c4.s42_lod_title)}</p>")
    s.append(tbl(c4.s42_lod_table.headers or ["名称","保留时间(min)","峰面积","S/N","浓度(ng/ml)","相当于样品的量(ppm)"],
                 c4.s42_lod_table.rows or []))
    s.append(f"<p>结论：{e(c4.s42_conclusion)}</p>")

    s.append("<h3>4.3 线性与范围</h3>")
    for i, name in enumerate(names):
        title = c4.s43_titles[i] if i < len(c4.s43_titles) else f"表4.3-{i+1} {name}线性"
        lt = c4.s43_tables[i] if i < len(c4.s43_tables) else None
        col_h = lt.col_headers if lt and lt.col_headers else ["定量限均值","S1","S2","S3","S4","S5"]
        span = max(len(col_h), 1)
        eq_val = lt.equation if lt else ""
        r2_val = lt.r2 if lt else ""
        s.append(f"<p>{e(title)}</p>")
        ths = "".join(f"<th>{e(str(h))}</th>" for h in ["名称"] + col_h)
        conc_row = "<tr><td>浓度(ng/ml)</td>" + "".join(f"<td>{e(str(v))}</td>" for v in (lt.concentrations if lt else [])) + "</tr>"
        area_row = "<tr><td>峰面积</td>" + "".join(f"<td>{e(str(v))}</td>" for v in (lt.peak_areas if lt else [])) + "</tr>"
        eq_row = f"<tr><td>线性方程</td><td colspan=\"{span}\">{e(eq_val)}</td></tr>"
        r_row = f"<tr><td>相关系数</td><td colspan=\"{span}\">{e(r2_val)}</td></tr>"
        s.append(f"<table><thead><tr>{ths}</tr></thead><tbody>{conc_row}{area_row}{eq_row}{r_row}</tbody></table>")
    s.append(f"<p>结论：{e(c4.s43_conclusion)}</p>")

    s.append("<h3>4.4 准确度和重复性</h3>")
    for i, name in enumerate(names):
        title = c4.s44_titles[i] if i < len(c4.s44_titles) else f"表4.4-{i+1} {name}回收率"
        td = c4.s44_tables[i] if i < len(c4.s44_tables) else None
        s.append(f"<p>{e(title)}</p>")
        if td and getattr(td, 'sub_rows', None) is not None:
            s.append(recovery_tbl(td))
        else:
            s.append(tbl(td.headers if td and td.headers else ["NDMA","加入量(ng/ml)","样本含量(ng/ml)","测得量(ng/ml)","回收率(%)","平均回收率(%)","RSD(%)"],
                         td.rows if td else []))
    s.append(f"<p>结论：{e(c4.s44_conclusion)}</p>")

    s.append(f"<h3>4.5 样品检测</h3><p>{e(c4.s45_table_title)}</p>")
    s45h = c4.s45_table.headers or (["批号", "称样量(mg)"] + [v for _ in names for v in ["检测值(ppm)", "均值(ppm)"]])
    s45rows = c4.s45_table.rows or []
    groups = [list(range(i, min(i + 2, len(names)))) for i in range(0, len(names), 2)] or [[0]]
    for gi, group in enumerate(groups):
        if len(groups) > 1:
            s.append(f"<p>检测结果表 {gi + 1}</p>")
        table_html = ["<table><thead>"]
        table_html.append("<tr>")
        table_html.append(f"<th rowspan=\"2\">{e(s45h[0] if len(s45h) > 0 else '批号')}</th>")
        table_html.append(f"<th rowspan=\"2\">{e(s45h[1] if len(s45h) > 1 else '称样量(mg)')}</th>")
        for compound_idx in group:
            nm = names[compound_idx] if compound_idx < len(names) and names[compound_idx] else f"对照品{compound_idx + 1}"
            table_html.append(f"<th colspan=\"2\">{e(nm)}</th>")
        table_html.append("</tr><tr>")
        for compound_idx in group:
            table_html.append(f"<th>{e(s45h[2 + compound_idx * 2] if 2 + compound_idx * 2 < len(s45h) else '检测值(ppm)')}</th>")
            table_html.append(f"<th>{e(s45h[3 + compound_idx * 2] if 3 + compound_idx * 2 < len(s45h) else '均值(ppm)')}</th>")
        table_html.append("</tr></thead><tbody>")
        even_rows = list(s45rows)
        if len(even_rows) % 2 != 0:
            even_rows.append([""] * max(len(s45h), 2 + len(names) * 2))
        if not even_rows:
            even_rows = [[""] * max(len(s45h), 2 + len(names) * 2) for _ in range(2)]
        for pair in range(len(even_rows) // 2):
            row0 = even_rows[pair * 2]
            row1 = even_rows[pair * 2 + 1]
            table_html.append("<tr>")
            table_html.append(f"<td rowspan=\"2\">{e(str(row0[0] if len(row0) > 0 else ''))}</td>")
            table_html.append(f"<td>{e(str(row0[1] if len(row0) > 1 else ''))}</td>")
            for compound_idx in group:
                det_idx = 2 + compound_idx * 2
                avg_idx = det_idx + 1
                table_html.append(f"<td>{e(str(row0[det_idx] if det_idx < len(row0) else ''))}</td>")
                table_html.append(f"<td rowspan=\"2\">{e(str(row0[avg_idx] if avg_idx < len(row0) else ''))}</td>")
            table_html.append("</tr><tr>")
            table_html.append(f"<td>{e(str(row1[1] if len(row1) > 1 else ''))}</td>")
            for compound_idx in group:
                det_idx = 2 + compound_idx * 2
                table_html.append(f"<td>{e(str(row1[det_idx] if det_idx < len(row1) else ''))}</td>")
            table_html.append("</tr>")
        table_html.append("</tbody></table>")
        s.append("".join(table_html))
    s.append(f"<p>结论：{e(c4.s45_conclusion)}</p>")

    s.append("<h2>5. 修订历史</h2>")
    s.append(tbl(["文件编号","更改原因","生效日期"], p.revision_rows or []))

    body = "\n".join(s)
    return f"""<!DOCTYPE html>
<html lang="zh"><head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{e(p.report_title or '报告预览')}</title>
<style>
body{{font-family:'宋体',serif;font-size:14px;margin:16px;line-height:1.6;color:#222}}
h1{{text-align:center;font-size:20px}}h2{{font-size:15px;border-bottom:1px solid #ccc;margin-top:20px}}
h3{{font-size:13px;margin-top:12px}}
table{{border-collapse:collapse;width:100%;margin-bottom:12px;font-size:12px}}
th,td{{border:1px solid #bbb;padding:4px 8px;text-align:left}}
th{{background:#f0f4fa;font-weight:bold}}p{{margin:4px 0}}
</style></head><body>{body}</body></html>"""


# ── PDF 生成 ──────────────────────────────────────────────────────────────────

def render_pdf(output_path: str, payload) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import glob as _glob, os, xml.sax.saxutils as _sax
    from datetime import datetime as _dt

    font_name = "Helvetica"
    env_font = os.environ.get("CJK_FONT_PATH")
    for pat in ([env_font] if env_font else []) + [
        "C:/Windows/Fonts/simsun.ttc", "C:/Windows/Fonts/msyh.ttc", "C:/Windows/Fonts/simhei.ttf",
        "/system/fonts/NotoSansCJK-Regular.ttc", "/system/fonts/DroidSansFallback.ttf",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttf", "/System/Library/Fonts/PingFang.ttc",
    ]:
        found = _glob.glob(pat)
        if found:
            try:
                pdfmetrics.registerFont(TTFont("CJK", found[0]))
                font_name = "CJK"
                break
            except Exception:
                continue

    h1s = ParagraphStyle("h1", fontName=font_name, fontSize=13, spaceBefore=12, spaceAfter=4, leading=18)
    h2s = ParagraphStyle("h2", fontName=font_name, fontSize=11, spaceBefore=8, spaceAfter=3, leading=16)
    bs  = ParagraphStyle("b",  fontName=font_name, fontSize=10, spaceAfter=3, leading=14)
    cvr = ParagraphStyle("cv", fontName=font_name, fontSize=18, alignment=1, spaceAfter=8, leading=24)
    cs  = ParagraphStyle("c",  fontName=font_name, fontSize=9, leading=12, wordWrap='CJK')

    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm, topMargin=2.5*cm, bottomMargin=2.5*cm)

    def _c(v): return Paragraph(_sax.escape(str(v or "")), cs)
    def _p(v, st): return Paragraph(_sax.escape(str(v or "")), st)

    def tbl(headers, rows):
        if not headers:
            return []
        data = [[_c(h) for h in headers]] + ([[_c(v) for v in r] for r in rows] if rows else [[_c("") for _ in headers]])
        col_w = doc.width / max(len(headers), 1)
        t = Table(data, colWidths=[col_w]*len(headers), repeatRows=1)
        t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#e8edf5")),
                                ("GRID",(0,0),(-1,-1),0.5,colors.grey),("VALIGN",(0,0),(-1,-1),"TOP")]))
        return [t, Spacer(1, 6)]

    def recovery_tbl(tbl_def):
        headers = _build_recovery_headers(tbl_def)
        rows = _build_recovery_rows(tbl_def)
        if not rows:
            rows = [["", "", "", "", "", "", ""] for _ in range(6)]
        data = [[_c(h) for h in headers]] + [[_c("") for _ in headers] for _ in rows]
        for i, row in enumerate(rows, start=1):
            for col_idx in range(min(5, len(row))):
                data[i][col_idx] = _c(row[col_idx])
            if len(row) > 5:
                data[i][5] = _c(row[5] if i == 1 else '')
            if len(row) > 6:
                data[i][6] = _c(row[6] if i == 1 else '')
        col_w = doc.width / max(len(headers), 1)
        t = Table(data, colWidths=[col_w]*len(headers), repeatRows=1)
        style_cmds = [("BACKGROUND",(0,0),(-1,0),colors.HexColor("#e8edf5")),
                      ("GRID",(0,0),(-1,-1),0.5,colors.grey),
                      ("VALIGN",(0,0),(-1,-1),"TOP")]
        if len(rows) > 1:
            for col in [5, 6]:
                style_cmds.append(("SPAN", (col, 1), (col, len(rows))))
        t.setStyle(TableStyle(style_cmds))
        return [t, Spacer(1, 6)]

    def _build_s45_pdf_story(story, tbl_def, names, _c, _p, bs):
        headers = (tbl_def.headers if tbl_def and getattr(tbl_def, 'headers', None) else [])
        rows = list(tbl_def.rows if tbl_def and tbl_def.rows else [])
        if not names:
            names = [f"对照品{i+1}" for i in range(max(1, (max(len(headers), 2) - 2) // 2))]
        expected_cols = 2 + len(names) * 2
        if len(rows) % 2 != 0:
            rows.append([""] * expected_cols)
        if not rows:
            rows = [[""] * expected_cols for _ in range(2)]

        def _header(idx, fallback):
            return headers[idx] if idx < len(headers) and headers[idx] else fallback

        groups = [list(range(i, min(i + 2, len(names)))) for i in range(0, len(names), 2)] or [[0]]
        for gi, group in enumerate(groups):
            if len(groups) > 1:
                story.append(_p(f"检测结果表 {gi + 1}", bs))
            data = []
            top = [_c(_header(0, "批号")), _c(_header(1, "称样量(mg)"))]
            sub = [_c(""), _c("")]
            for compound_idx in group:
                nm = names[compound_idx] if compound_idx < len(names) and names[compound_idx] else f"对照品{compound_idx + 1}"
                top.extend([_c(nm), _c("")])
                sub.extend([
                    _c(_header(2 + compound_idx * 2, "检测值(ppm)")),
                    _c(_header(3 + compound_idx * 2, "均值(ppm)")),
                ])
            data.extend([top, sub])
            spans = [((0, 0), (0, 1)), ((1, 0), (1, 1))]
            for local_idx in range(len(group)):
                base = 2 + local_idx * 2
                spans.append(((base, 0), (base + 1, 0)))
            for pair in range(len(rows) // 2):
                row0 = rows[pair * 2]
                row1 = rows[pair * 2 + 1]
                r0 = len(data)
                r1 = r0 + 1
                first = [_c(row0[0] if len(row0) > 0 else ""), _c(row0[1] if len(row0) > 1 else "")]
                second = [_c(""), _c(row1[1] if len(row1) > 1 else "")]
                spans.append(((0, r0), (0, r1)))
                for local_idx, compound_idx in enumerate(group):
                    source_base = 2 + compound_idx * 2
                    first.extend([
                        _c(row0[source_base] if source_base < len(row0) else ""),
                        _c(row0[source_base + 1] if source_base + 1 < len(row0) else ""),
                    ])
                    second.extend([
                        _c(row1[source_base] if source_base < len(row1) else ""),
                        _c(""),
                    ])
                    avg_col = 2 + local_idx * 2 + 1
                    spans.append(((avg_col, r0), (avg_col, r1)))
                data.extend([first, second])
            col_widths = [doc.width * 0.14, doc.width * 0.12] + [doc.width * 0.14, doc.width * 0.12] * len(group)
            t = Table(data, colWidths=col_widths, repeatRows=2)
            style_cmds = [
                ("BACKGROUND", (0, 0), (-1, 1), colors.HexColor("#e8edf5")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
            for start, end in spans:
                style_cmds.append(("SPAN", start, end))
            t.setStyle(TableStyle(style_cmds))
            story.extend([t, Spacer(1, 6)])

    p = payload
    c4 = p.chapter4
    names = [c.name for c in (c4.compounds or [])]
    story = []

    story += [_p(p.report_title or "化学检测实验报告", cvr),
              _p(f"文档编号：{p.doc_number}  密级：{p.security_level}  项目名称：{p.project_name}", bs),
              _p(f"生成日期：{_dt.now().strftime('%Y-%m-%d')}", bs), Spacer(1,8)]
    story += tbl(["编制人","审核人","批准人","编制日期","版本号"],
                 [[r.editor,r.reviewer,r.approver,r.date,r.version] for r in (p.personnel or [])])

    story += [_p("1. 检测方法", h1s), _p(p.detection_method, bs)]
    story += tbl(["杂质名称","每日最大允许摄入量","每日最大摄入量","限度","供试品浓度","限度浓度"], p.impurities or [])
    story += [_p("2. 方法学验证标准", h1s)]
    story += tbl(["项目","可接受标准","结果"], p.validation_rows or [])
    story += [_p("3. 仪器与试剂", h1s)]
    story += tbl(["名称","型号","编号","厂家","有效期"], p.instruments or [])
    story += tbl(["名称","批号","级别","厂家"], p.reagents or [])
    story += tbl(["名称","批号","含量（%）","厂家"], p.references or [])

    story += [_p("4. 验证内容", h1s), _p("4.1 专属性", h2s), _p(c4.s41_table_title, bs)]
    h41 = c4.s41_table.headers or ["杂质名称","空白溶剂保留时间min","限度定位溶液保留时间min","供试品溶液保留时间min"]
    story += tbl(h41, c4.s41_table.rows or [])
    story += [_p(f"结论：{c4.s41_conclusion}", bs), _p("4.2 定量限和检测限", h2s)]
    for i, name in enumerate(names):
        td = c4.s42_loq_tables[i] if i < len(c4.s42_loq_tables) else None
        story += [_p(c4.s42_loq_titles[i] if i < len(c4.s42_loq_titles) else f"表4.2-{i+1} {name}定量限", bs)]
        story += tbl(td.headers if td and td.headers else ["序号","保留时间(min)","峰面积","S/N","浓度(ng/ml)","相当于样品的量(ppm)"],
                     td.rows if td else [])
    story += [_p(c4.s42_lod_title, bs)]
    story += tbl(c4.s42_lod_table.headers or ["名称","保留时间(min)","峰面积","S/N","浓度(ng/ml)","相当于样品的量(ppm)"],
                 c4.s42_lod_table.rows or [])
    story += [_p(f"结论：{c4.s42_conclusion}", bs), _p("4.3 线性与范围", h2s)]
    for i, name in enumerate(names):
        lt = c4.s43_tables[i] if i < len(c4.s43_tables) else None
        col_h = lt.col_headers if lt and lt.col_headers else ["定量限均值","S1","S2","S3","S4","S5"]
        story += [_p(c4.s43_titles[i] if i < len(c4.s43_titles) else f"表4.3-{i+1} {name}线性", bs)]
        span = len(col_h)
        eq_val = lt.equation if lt else ""
        r2_val = lt.r2 if lt else ""
        data = (
            [["名称"] + col_h]
            + [["浓度(ng/ml)"] + list(lt.concentrations if lt else [])]
            + [["峰面积"] + list(lt.peak_areas if lt else [])]
            + [["线性方程", eq_val] + [""] * (span - 1)]
            + [["相关系数", r2_val] + [""] * (span - 1)]
        )
        col_w = doc.width / max(len(col_h) + 1, 1)
        t = Table([[_c(v) for v in row] for row in data], colWidths=[col_w] * (len(col_h) + 1), repeatRows=1)
        style_cmds = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8edf5")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("SPAN", (1, 3), (-1, 3)),
            ("SPAN", (1, 4), (-1, 4)),
        ]
        t.setStyle(TableStyle(style_cmds))
        story += [t, Spacer(1, 6)]
    story += [_p(f"结论：{c4.s43_conclusion}", bs), _p("4.4 准确度和重复性", h2s)]
    for i, name in enumerate(names):
        td = c4.s44_tables[i] if i < len(c4.s44_tables) else None
        story += [_p(c4.s44_titles[i] if i < len(c4.s44_titles) else f"表4.4-{i+1} {name}回收率", bs)]
        if td and getattr(td, 'sub_rows', None) is not None:
            story += recovery_tbl(td)
        else:
            story += tbl(td.headers if td and td.headers else ["NDMA","加入量(ng/ml)","样本含量(ng/ml)","测得量(ng/ml)","回收率(%)","平均回收率(%)","RSD(%)"],
                         td.rows if td else [])
    story += [_p(f"结论：{c4.s44_conclusion}", bs), _p("4.5 样品检测", h2s), _p(c4.s45_table_title, bs)]
    _build_s45_pdf_story(story, c4.s45_table, names, _c, _p, bs)
    story += [_p(f"结论：{c4.s45_conclusion}", bs), _p("5. 修订历史", h1s)]
    story += tbl(["文件编号","更改原因","生效日期"], p.revision_rows or [])

    def _footer(canvas, d):
        canvas.saveState()
        canvas.setFont(font_name, 9)
        canvas.drawCentredString(A4[0]/2, 1.2*cm, f"第 {d.page} 页")
        canvas.restoreState()

    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)

