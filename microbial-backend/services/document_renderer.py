from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import os, tempfile, base64, html as html_module
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

UPLOAD_DIR = Path(os.environ.get("UPLOAD_DIR", "static/uploads"))
REPORTS_DIR = Path(os.environ.get("REPORTS_DIR", "static/reports"))

_SONG = "宋体"
_TNR = "Times New Roman"
_PT_BODY = Pt(12)
_PT_TABLE = Pt(10.5)
_PT_H1 = Pt(12)


def _set_run_font(run, bold=False, size=None):
    run.bold = bold
    if size:
        run.font.size = size
    run.font.name = _TNR
    run._element.rPr.rFonts.set(qn("w:eastAsia"), _SONG)


def _add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, bold=True, size=_PT_H1)


def _add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, bold=True, size=_PT_BODY)


def _set_cell_text(cell, text, bold=False):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    cell.paragraphs[0].clear()
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = cell.paragraphs[0].add_run(str(text))
    _set_run_font(run, bold=bold, size=_PT_TABLE)


def _add_table(doc, headers, rows):
    if not headers:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row[:len(headers)]):
            _set_cell_text(table.rows[r_idx + 1].cells[c_idx], val)
    doc.add_paragraph()


def _find_upload_file(file_id):
    if not file_id:
        return None
    path = UPLOAD_DIR / file_id
    return path if path.exists() else None


def _resize_image(image_path, max_width_inches=4.0):
    img = Image.open(image_path)
    dpi = img.info.get("dpi", (96, 96))[0]
    max_px = int(max_width_inches * dpi)
    if img.width > max_px:
        ratio = max_px / img.width
        img = img.resize((max_px, int(img.height * ratio)), Image.LANCZOS)
    out = image_path + "_resized.png"
    img.save(out, "PNG")
    return out


def _insert_signature_into_cell(cell, file_id, width_inches=1.4):
    img_path = _find_upload_file(file_id)
    if not img_path:
        return False
    resized = _resize_image(str(img_path), width_inches)
    try:
        cell.paragraphs[0].clear()
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].add_run().add_picture(resized, width=Inches(width_inches))
        return True
    finally:
        if os.path.exists(resized):
            os.remove(resized)


def _signature_img_html(file_id, max_width=140, max_height=60):
    img_path = _find_upload_file(file_id)
    if not img_path:
        return ""
    try:
        b64 = base64.b64encode(img_path.read_bytes()).decode()
        suffix = img_path.suffix.lower().lstrip('.')
        mime = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png", "bmp": "bmp", "tiff": "tiff"}.get(suffix, "png")
        return f'<img src="data:image/{mime};base64,{b64}" style="max-width:{max_width}px;max-height:{max_height}px">'
    except Exception:
        return ""


def _plot_linearity(xs, ys, equation, r2):
    plt.rcParams['font.sans-serif'] = ['SimSun', 'Microsoft YaHei', 'DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False
    fig, ax = plt.subplots(figsize=(5, 3.5))
    ax.scatter(xs, ys, color='steelblue', zorder=3)
    m = np.polyfit(xs, ys, 1)
    x_line = np.linspace(min(xs), max(xs), 100)
    ax.plot(x_line, np.polyval(m, x_line), 'r--', linewidth=1)
    ax.set_xlabel('稀释度/浓度', fontsize=9)
    ax.set_ylabel('OD600/菌落数', fontsize=9)
    ax.grid(True, linestyle='--', alpha=0.4)
    ax.text(0.97, 0.05, f'{equation}\nR²={r2:.4f}',
            transform=ax.transAxes, fontsize=8, ha='right', va='bottom',
            bbox=dict(boxstyle='round,pad=0.3', facecolor='lightyellow', alpha=0.8))
    plt.tight_layout()
    tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
    fig.savefig(tmp.name, dpi=150)
    plt.close(fig)
    return tmp.name


def _add_page_numbers(doc, first_page_footer=""):
    def _make_fld(t):
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), t)
        return fld

    def _make_instr(text):
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = text
        return instr

    def _page_num_run(para):
        r = OxmlElement('w:r')
        r.append(_make_fld('begin'))
        r.append(_make_instr(' PAGE '))
        r.append(_make_fld('separate'))
        r.append(_make_fld('end'))
        para._p.append(r)

    section = doc.sections[0]
    section.different_first_page_header_footer = True
    fp = section.first_page_footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(fp.add_run(first_page_footer or ""), size=Pt(10))

    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p.add_run("第 "), size=Pt(10))
    _page_num_run(p)
    _set_run_font(p.add_run(" 页"), size=Pt(10))


def _render_qc_section(doc, payload):
    if payload.qc_concentrations or payload.qc_od_values:
        _add_h2(doc, "附表1 数据处理")
        headers = payload.qc_data_headers or []
        conc_row = ["稀释度/浓度"] + list(payload.qc_concentrations or [])
        od_row = ["OD600/菌落数"] + list(payload.qc_od_values or [])
        if headers:
            _add_table(doc, headers, [conc_row, od_row])
        if payload.qc_linear_equation:
            _set_run_font(doc.add_paragraph().add_run(f"线性方程：{payload.qc_linear_equation}"), size=_PT_BODY)
        if payload.qc_correlation_coeff is not None:
            _set_run_font(doc.add_paragraph().add_run(f"相关系数 R：{payload.qc_correlation_coeff:.6f}"), size=_PT_BODY)

    if payload.qc_conclusion_headers and payload.qc_conclusion_rows:
        doc.add_page_break()
        _add_h2(doc, "附表2 结论")
        try:
            xs = [float(v) for v in (payload.qc_concentrations or []) if v]
            ys = [float(v) for v in (payload.qc_od_values or []) if v]
            if len(xs) >= 2 and len(xs) == len(ys):
                chart_path = _plot_linearity(xs, ys, payload.qc_linear_equation or '', payload.qc_correlation_coeff or 0.0)
                doc.add_picture(chart_path, width=Inches(4.5))
                os.remove(chart_path)
        except Exception:
            pass
        row_labels = payload.qc_conclusion_row_labels or ['稀释度/浓度', 'OD600/菌落数']
        n_data = len(payload.qc_conclusion_headers) - 1
        v0 = payload.qc_conclusion_rows[0] if payload.qc_conclusion_rows else []
        v1 = payload.qc_conclusion_rows[1] if len(payload.qc_conclusion_rows) > 1 else []
        _add_table(doc, payload.qc_conclusion_headers, [
            [row_labels[0]] + list(v0[:n_data]),
            [row_labels[1] if len(row_labels) > 1 else 'OD600/菌落数'] + list(v1[:n_data]),
        ])


def render_report(output_path, payload):
    doc = Document()

    # 封面
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(title_para.add_run(payload.report_title or "微生物实验报告"), bold=True, size=Pt(22))

    for label, val in [
        ("文档编号", payload.doc_number),
        ("密级", payload.security_level),
        ("项目名称", payload.project_name),
        ("生成日期", datetime.now().strftime("%Y-%m-%d")),
    ]:
        _set_run_font(doc.add_paragraph().add_run(f"{label}：{val}"), size=_PT_BODY)
    doc.add_paragraph()

    _add_h2(doc, "实验人员信息")
    _add_table(doc, ["编制人", "审核人", "批准人", "编制日期", "版本号"],
               [[p.editor, p.reviewer, ('[签名]' if idx < len(payload.approver_signatures or []) and (payload.approver_signatures or [])[idx] else p.approver), p.date, p.version] for idx, p in enumerate(payload.personnel)])
    for idx, signature in enumerate(payload.approver_signatures or []):
        if idx < len(payload.personnel):
            _insert_signature_into_cell(doc.tables[-1].rows[idx + 1].cells[2], signature)

    # 一、材料试剂与仪器
    _add_h1(doc, "一、材料试剂与仪器")
    _add_h2(doc, "1. 菌株信息")
    _add_table(doc, payload.strains_headers or ["菌株名称", "来源", "保存日期", "保存方式"], payload.strains)
    _add_h2(doc, "2. 培养基")
    _add_table(doc, payload.media_headers or ["名称", "成分", "配制日期", "灭菌方式"], payload.media)
    _add_h2(doc, "3. 试剂信息")
    _add_table(doc, payload.reagents_headers or ["名称", "厂家", "批号"], payload.reagents)
    _add_h2(doc, "4. 仪器信息")
    _add_table(doc, payload.instruments_headers or ["名称", "厂家", "编号", "型号", "校准日期"], payload.instruments)
    _add_h2(doc, "5. 耗材清单")
    _add_table(doc, payload.consumables_headers or ["名称", "厂家", "生产日期", "有效期"], payload.consumables)

    # 二、操作步骤
    _add_h1(doc, "二、操作步骤")
    fixed_titles = ["1. 菌株活化与培养", "2. 接种与发酵", "3. 样品处理与检测"]
    default_hdrs = ["序号", "步骤名称", "具体步骤"]
    proc_headers = payload.procedure_headers or []
    for i, title in enumerate(fixed_titles):
        _add_h2(doc, title)
        hdrs = proc_headers[i] if i < len(proc_headers) else default_hdrs
        rows = payload.procedures[i] if i < len(payload.procedures) else []
        _add_table(doc, hdrs, rows)
    for i, dyn_title in enumerate(payload.dynamic_section_titles or []):
        _add_h2(doc, f"{len(fixed_titles) + i + 1}. {dyn_title}")
        hdrs = payload.dynamic_section_headers[i] if i < len(payload.dynamic_section_headers) else default_hdrs
        rows = payload.dynamic_section_tables[i] if i < len(payload.dynamic_section_tables) else []
        _add_table(doc, hdrs, rows)

    # 三、实验结果与分析
    _add_h1(doc, "三、实验结果与分析")
    result_headers = ["序号", "步骤名称", "图谱"] + list(payload.result_extra_headers or [])
    result_table = doc.add_table(rows=1 + len(payload.results), cols=len(result_headers))
    result_table.style = "Table Grid"
    for i, h in enumerate(result_headers):
        _set_cell_text(result_table.rows[0].cells[i], h, bold=True)
    for r_idx, row in enumerate(payload.results):
        _set_cell_text(result_table.rows[r_idx + 1].cells[0], row.seq)
        _set_cell_text(result_table.rows[r_idx + 1].cells[1], row.name)
        for extra_idx, val in enumerate(row.extra_cols or []):
            cell_idx = 3 + extra_idx
            if cell_idx < len(result_headers):
                _set_cell_text(result_table.rows[r_idx + 1].cells[cell_idx], val)
        if row.image_file_id:
            img_path = UPLOAD_DIR / row.image_file_id
            if img_path.exists():
                resized = _resize_image(str(img_path), 2.5)
                result_table.rows[r_idx + 1].cells[2].paragraphs[0].add_run().add_picture(resized, width=Inches(2.5))
                if os.path.exists(resized):
                    os.remove(resized)
    doc.add_paragraph()

    # 四、实验结论
    _add_h1(doc, "四、实验结论")

    # 渲染来自Flutter Step4的通用实验结果（表格或文本模式）
    if payload.results_mode == 'text':
        p = doc.add_paragraph(payload.results_text_content or '')
        if p.runs:
            _set_run_font(p.runs[0], size=Pt(10.5))
    elif payload.results_headers:
        _add_table(doc, payload.results_headers, payload.results_rows or [])
    else:
        _add_table(doc, payload.conclusion_table_headers or ["实验结论", "备注"],
                   payload.conclusion_table_rows or [])

    # 附录
    doc.add_page_break()
    _add_h1(doc, "附录")
    _render_qc_section(doc, payload)
    for fpath in (payload.appendix_file_paths or []):
        p_path = Path(fpath) if os.path.isabs(fpath) else UPLOAD_DIR / Path(fpath).name
        if p_path.exists() and p_path.suffix.lower() == '.txt':
            _add_h2(doc, f"附件：{p_path.name}")
            try:
                _set_run_font(doc.add_paragraph().add_run(p_path.read_text(encoding='utf-8', errors='replace')[:5000]), size=Pt(10))
            except Exception:
                pass

    _add_page_numbers(doc, payload.first_page_footer)
    doc.save(output_path)


def render_preview_html(payload) -> str:
    e = html_module.escape

    def table_html(headers, rows):
        if not headers:
            return ""
        ths = "".join(f"<th>{e(h)}</th>" for h in headers)
        trs = "".join(
            "<tr>" + "".join(f"<td>{e(str(v))}</td>" for v in row) + "</tr>"
            for row in rows
        )
        return f"<table><thead><tr>{ths}</tr></thead><tbody>{trs}</tbody></table>"

    s = []
    s.append(f"<h1 class='cover'>{e(payload.report_title or '微生物实验报告')}</h1>")
    s.append(f"<p>文档编号：{e(payload.doc_number)} &nbsp; 密级：{e(payload.security_level)}</p>")
    s.append(f"<p>项目名称：{e(payload.project_name)} &nbsp; 生成日期：{datetime.now().strftime('%Y-%m-%d')}</p>")
    s.append("<h2>实验人员信息</h2>")
    personnel_rows = []
    signatures = payload.approver_signatures or []
    for idx, p in enumerate(payload.personnel):
        personnel_rows.append([
            p.editor,
            p.reviewer,
            _signature_img_html(signatures[idx] if idx < len(signatures) else None) or p.approver,
            p.date,
            p.version,
        ])
    s.append(table_html(["编制人", "审核人", "批准人", "编制日期", "版本号"], personnel_rows))

    s.append("<h2>一、材料试剂与仪器</h2>")
    for title, hdrs, rows in [
        ("1. 菌株信息", payload.strains_headers or ["菌株名称", "来源", "保存日期", "保存方式"], payload.strains),
        ("2. 培养基", payload.media_headers or ["名称", "成分", "配制日期", "灭菌方式"], payload.media),
        ("3. 试剂信息", payload.reagents_headers or ["名称", "厂家", "批号"], payload.reagents),
        ("4. 仪器信息", payload.instruments_headers or ["名称", "厂家", "编号", "校准日期"], payload.instruments),
        ("5. 耗材清单", payload.consumables_headers or ["名称", "厂家", "生产日期", "有效期"], payload.consumables),
    ]:
        s.append(f"<h3>{e(title)}</h3>")
        s.append(table_html(hdrs, rows))

    s.append("<h2>二、操作步骤</h2>")
    fixed_titles = ["1. 菌株活化与培养", "2. 接种与发酵", "3. 样品处理与检测"]
    proc_headers = payload.procedure_headers or []
    for i, title in enumerate(fixed_titles):
        s.append(f"<h3>{e(title)}</h3>")
        hdrs = proc_headers[i] if i < len(proc_headers) else ["序号", "步骤名称", "具体步骤"]
        s.append(table_html(hdrs, payload.procedures[i] if i < len(payload.procedures) else []))
    for i, dyn in enumerate(payload.dynamic_section_titles or []):
        s.append(f"<h3>{len(fixed_titles)+i+1}. {e(dyn)}</h3>")
        hdrs = payload.dynamic_section_headers[i] if i < len(payload.dynamic_section_headers) else ["序号", "步骤名称", "具体步骤"]
        s.append(table_html(hdrs, payload.dynamic_section_tables[i] if i < len(payload.dynamic_section_tables) else []))

    s.append("<h2>三、实验结果与分析</h2>")
    result_headers = ["序号", "步骤名称", "图谱"] + list(payload.result_extra_headers or [])
    ths = "".join(f"<th>{e(h)}</th>" for h in result_headers)
    trs = ""
    for r in payload.results:
        img_cell = ""
        if r.image_file_id:
            img_path = UPLOAD_DIR / r.image_file_id
            if img_path.exists():
                try:
                    b64 = base64.b64encode(img_path.read_bytes()).decode()
                    suffix = img_path.suffix.lower().lstrip(".")
                    mime = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png", "bmp": "bmp"}.get(suffix, "png")
                    img_cell = f'<img src="data:image/{mime};base64,{b64}" style="max-width:200px;max-height:160px">'
                except Exception:
                    img_cell = "(图谱)"
        extra_tds = "".join(f"<td>{e(str(v))}</td>" for v in (r.extra_cols or []))
        trs += f"<tr><td>{e(r.seq)}</td><td>{e(r.name)}</td><td>{img_cell}</td>{extra_tds}</tr>"
    s.append(f"<table><thead><tr>{ths}</tr></thead><tbody>{trs}</tbody></table>")

    s.append("<h2>四、实验结论</h2>")
    if payload.results_mode == 'text':
        s.append(f"<p>{e(payload.results_text_content or '')}</p>")
    elif payload.results_headers:
        s.append(table_html(payload.results_headers, payload.results_rows or []))
    else:
        s.append(table_html(payload.conclusion_table_headers or ["实验结论", "备注"], payload.conclusion_table_rows or []))

    s.append("<hr><h2>附录</h2>")
    if payload.qc_concentrations or payload.qc_od_values:
        s.append("<h3>数据处理</h3>")
        if payload.qc_data_headers:
            s.append(table_html(payload.qc_data_headers, [
                ["稀释度/浓度"] + list(payload.qc_concentrations or []),
                ["OD600/菌落数"] + list(payload.qc_od_values or []),
            ]))
        if payload.qc_linear_equation:
            s.append(f"<p>线性方程：{e(payload.qc_linear_equation)}</p>")
        if payload.qc_correlation_coeff is not None:
            s.append(f"<p>相关系数 R：{payload.qc_correlation_coeff:.6f}</p>")
    if payload.qc_conclusion_headers and payload.qc_conclusion_rows:
        s.append("<h3>结论</h3>")
        try:
            xs = [float(v) for v in (payload.qc_concentrations or []) if v]
            ys = [float(v) for v in (payload.qc_od_values or []) if v]
            if len(xs) >= 2 and len(xs) == len(ys):
                chart_path = _plot_linearity(xs, ys, payload.qc_linear_equation or '', payload.qc_correlation_coeff or 0.0)
                b64 = base64.b64encode(Path(chart_path).read_bytes()).decode()
                os.remove(chart_path)
                s.append(f'<img src="data:image/png;base64,{b64}" style="max-width:400px;display:block;margin:8px 0">')
        except Exception:
            pass
        rl = payload.qc_conclusion_row_labels or ['稀释度/浓度', 'OD600/菌落数']
        n_data = len(payload.qc_conclusion_headers) - 1
        v0 = payload.qc_conclusion_rows[0] if payload.qc_conclusion_rows else []
        v1 = payload.qc_conclusion_rows[1] if len(payload.qc_conclusion_rows) > 1 else []
        s.append(table_html(payload.qc_conclusion_headers, [
            [rl[0]] + list(v0[:n_data]),
            [rl[1] if len(rl) > 1 else 'OD600/菌落数'] + list(v1[:n_data]),
        ]))
    for fpath in (payload.appendix_file_paths or []):
        p_path = Path(fpath) if os.path.isabs(fpath) else UPLOAD_DIR / Path(fpath).name
        if p_path.exists() and p_path.suffix.lower() == '.txt':
            try:
                content = p_path.read_text(encoding='utf-8', errors='replace')
                s.append(f"<h3>附件：{e(p_path.name)}</h3>")
                s.append(f"<pre style='white-space:pre-wrap;font-size:12px;background:#f9f9f9;padding:8px;border:1px solid #ddd'>{e(content[:5000])}</pre>")
            except Exception:
                pass

    body = "\n".join(s)
    return f"""<!DOCTYPE html>
<html lang="zh"><head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{e(payload.report_title or '报告预览')}</title>
<style>
body{{font-family:'宋体',serif;font-size:14px;margin:16px;line-height:1.6;color:#222}}
h1.cover{{text-align:center;font-size:22px}}
h2{{font-size:15px;border-bottom:1px solid #ccc;padding-bottom:2px;margin-top:20px}}
h3{{font-size:14px;margin-top:12px}}
table{{border-collapse:collapse;width:100%;margin-bottom:12px;font-size:13px}}
th,td{{border:1px solid #bbb;padding:4px 8px;text-align:left}}
th{{background:#f5f5f5;font-weight:bold}}
hr{{margin:24px 0}}
p{{margin:4px 0}}
</style>
</head><body>
{body}
</body></html>"""


def render_pdf(output_path, payload):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import glob as _glob
    import xml.sax.saxutils as _sax

    font_name = "Helvetica"
    for pattern in [
        os.environ.get("CJK_FONT_PATH", ""),
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/msyh.ttc",
        "/system/fonts/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttf",
    ]:
        if not pattern:
            continue
        found = _glob.glob(pattern)
        if found:
            try:
                pdfmetrics.registerFont(TTFont("CJK", found[0]))
                font_name = "CJK"
                break
            except Exception:
                continue

    cell_style = ParagraphStyle("cell", fontName=font_name, fontSize=9, leading=12, wordWrap='CJK')
    h1 = ParagraphStyle("h1", fontName=font_name, fontSize=13, spaceAfter=6, spaceBefore=12, leading=18)
    h2 = ParagraphStyle("h2", fontName=font_name, fontSize=11, spaceAfter=4, spaceBefore=8, leading=16)
    body = ParagraphStyle("body", fontName=font_name, fontSize=10, spaceAfter=3, leading=14)
    cover = ParagraphStyle("cover", fontName=font_name, fontSize=18, alignment=1, spaceAfter=8, leading=24)

    doc = SimpleDocTemplate(output_path, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2.5*cm, bottomMargin=2.5*cm)

    def _cell(v):
        return Paragraph(_sax.escape(str(v)), cell_style)

    def tbl(headers, rows):
        if not headers:
            return []
        data = [[_cell(h) for h in headers]] + ([[_cell(v) for v in row] for row in rows] if rows else [[_cell("") for _ in headers]])
        col_w = doc.width / max(len(headers), 1)
        t = Table(data, colWidths=[col_w] * len(headers), repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8edf5")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        return [t, Spacer(1, 6)]

    story = []
    story.append(Paragraph(_sax.escape(payload.report_title or "微生物实验报告"), cover))
    for label, val in [("文档编号", payload.doc_number), ("密级", payload.security_level),
                       ("项目名称", payload.project_name), ("生成日期", datetime.now().strftime("%Y-%m-%d"))]:
        story.append(Paragraph(f"{_sax.escape(label)}：{_sax.escape(str(val or ''))}", body))
    story.append(Spacer(1, 12))

    story.append(Paragraph("实验人员信息", h2))
    story.extend(tbl(["编制人", "审核人", "批准人", "编制日期", "版本号"],
                     [[p.editor, p.reviewer, p.approver, p.date, p.version] for p in payload.personnel]))

    story.append(Paragraph("一、材料试剂与仪器", h1))
    for title, hdrs, rows in [
        ("1. 菌株信息", payload.strains_headers or ["菌株名称", "来源", "保存日期", "保存方式"], payload.strains),
        ("2. 培养基", payload.media_headers or ["名称", "成分", "配制日期", "灭菌方式"], payload.media),
        ("3. 试剂信息", payload.reagents_headers or ["名称", "厂家", "批号"], payload.reagents),
        ("4. 仪器信息", payload.instruments_headers or ["名称", "厂家", "编号", "校准日期"], payload.instruments),
        ("5. 耗材清单", payload.consumables_headers or ["名称", "厂家", "生产日期", "有效期"], payload.consumables),
    ]:
        story.append(Paragraph(title, h2))
        story.extend(tbl(hdrs, rows))

    story.append(Paragraph("二、操作步骤", h1))
    fixed_titles = ["1. 菌株活化与培养", "2. 接种与发酵", "3. 样品处理与检测"]
    proc_hdrs = payload.procedure_headers or []
    for i, title in enumerate(fixed_titles):
        story.append(Paragraph(title, h2))
        hdrs = proc_hdrs[i] if i < len(proc_hdrs) else ["序号", "步骤名称", "具体步骤"]
        story.extend(tbl(hdrs, payload.procedures[i] if i < len(payload.procedures) else []))
    for i, dyn in enumerate(payload.dynamic_section_titles or []):
        story.append(Paragraph(f"{len(fixed_titles)+i+1}. {_sax.escape(dyn)}", h2))
        hdrs = payload.dynamic_section_headers[i] if i < len(payload.dynamic_section_headers) else ["序号", "步骤名称", "具体步骤"]
        story.extend(tbl(hdrs, payload.dynamic_section_tables[i] if i < len(payload.dynamic_section_tables) else []))

    story.append(Paragraph("三、实验结果与分析", h1))
    story.extend(tbl(["序号", "步骤名称"] + list(payload.result_extra_headers or []),
                     [[r.seq, r.name] + list(r.extra_cols or []) for r in payload.results]))

    story.append(Paragraph("四、实验结论", h1))
    if payload.results_mode == 'text':
        story.append(Paragraph(payload.results_text_content or '', body))
    elif payload.results_headers:
        story.extend(tbl(payload.results_headers, payload.results_rows or []))
    else:
        story.extend(tbl(payload.conclusion_table_headers or ["实验结论", "备注"], payload.conclusion_table_rows or []))

    story.append(PageBreak())
    story.append(Paragraph("附录", h1))

    footer_text = payload.first_page_footer or ""

    def on_first_page(canvas, doc):
        canvas.saveState()
        canvas.setFont(font_name, 9)
        canvas.drawCentredString(A4[0] / 2, 1.2 * cm, footer_text)
        canvas.restoreState()

    def on_later_pages(canvas, doc):
        canvas.saveState()
        canvas.setFont(font_name, 9)
        canvas.drawCentredString(A4[0] / 2, 1.2 * cm, f"第 {doc.page} 页")
        canvas.restoreState()

    doc.build(story, onFirstPage=on_first_page, onLaterPages=on_later_pages)
